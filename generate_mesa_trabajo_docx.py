# -*- coding: utf-8 -*-
"""Genera el documento Word de mesa de trabajo funcional del SGMC.

Salida: Definicion_Funcional_SGMC_Mesa_de_Trabajo.docx
Fuente de contenido: DEFINICION_FUNCIONAL_MESA_DE_TRABAJO.md
Figuras: Manuales/images/fig_0*.png (generadas por figuras.py)

El documento incluye casillas de respuesta editables para que el líder funcional
responda directamente sobre el archivo y lo devuelva.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(BASE, "Manuales", "images")
SALIDA = os.path.join(BASE, "Definicion_Funcional_SGMC_Mesa_de_Trabajo.docx")

TINTA = RGBColor(0x1C, 0x25, 0x30)
GRIS = RGBColor(0x6E, 0x7A, 0x86)
AZUL = RGBColor(0x1F, 0x5B, 0x99)
ROJO = RGBColor(0xA8, 0x2D, 0x2D)
AMBAR = RGBColor(0xB0, 0x70, 0x0C)
VERDE = RGBColor(0x20, 0x74, 0x54)


# --------------------------------------------------------------- utilidades
def sombrear(celda, hex_color):
    tc = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def bordes(celda, color="CBD3DB", sz=4, lados=("top", "left", "bottom", "right")):
    tc = celda._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for lado in lados:
        e = OxmlElement(f"w:{lado}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        b.append(e)
    tc.append(b)


def barra_izquierda(celda, color):
    tc = celda._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "18")
    e.set(qn("w:color"), color)
    b.append(e)
    tc.append(b)


def p(doc, texto="", size=10.5, bold=False, color=TINTA, space_after=6,
      space_before=0, align=None, italic=False, izq=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.left_indent = Cm(izq)
    if align:
        par.alignment = align
    if texto:
        r = par.add_run(texto)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        r.font.name = "Segoe UI"
    return par


def titulo(doc, texto, nivel=1, color=TINTA, size=None):
    tam = size or (16 if nivel == 1 else 13 if nivel == 2 else 11.5)
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(18 if nivel == 1 else 14)
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.keep_with_next = True
    r = par.add_run(texto)
    r.font.size = Pt(tam)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Segoe UI"
    return par


def vinneta(doc, texto, size=10.5, color=TINTA, izq=0.8):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.left_indent = Cm(izq)
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(texto)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Segoe UI"
    return par


def bloque(doc, etiqueta, texto, color_hex, color_txt, fondo_hex):
    """Bloque con barra de color a la izquierda."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, fondo_hex)
    barra_izquierda(c, color_hex)
    c.paragraphs[0].text = ""
    par = c.paragraphs[0]
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(6)
    if etiqueta:
        r = par.add_run(etiqueta + "  ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = color_txt
        r.font.name = "Segoe UI"
    r = par.add_run(texto)
    r.font.size = Pt(10)
    r.font.color.rgb = TINTA
    r.font.name = "Segoe UI"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def casilla_respuesta(doc, alto_lineas=5, etiqueta="RESPUESTA"):
    """Caja editable para que el funcional escriba."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, "FFFDF5")
    bordes(c, color="B0700C", sz=8)
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(etiqueta)
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = AMBAR
    r.font.name = "Segoe UI"
    for _ in range(alto_lineas):
        pp = c.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.space_before = Pt(0)
        rr = pp.add_run(" ")
        rr.font.size = Pt(10.5)
        rr.font.name = "Segoe UI"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def tabla(doc, encabezados, filas, anchos=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        sombrear(hdr[i], "1F5B99")
        par = hdr[i].paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.space_before = Pt(2)
        r = par.add_run(h)
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Segoe UI"
    for k, fila in enumerate(filas):
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            if k % 2 == 1:
                sombrear(celdas[i], "F4F6F8")
            par = celdas[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.space_before = Pt(2)
            r = par.add_run(str(v))
            r.font.size = Pt(size)
            r.font.name = "Segoe UI"
            r.font.color.rgb = TINTA
    if anchos:
        for i, w in enumerate(anchos):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def figura(doc, archivo, pie, ancho=16.5):
    ruta = os.path.join(IMG, archivo)
    if not os.path.exists(ruta):
        print("  AVISO: falta la figura", archivo)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(3)
    par.add_run().add_picture(ruta, width=Cm(ancho))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(pie)
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = GRIS
    r.font.name = "Segoe UI"


def pie_de_pagina(doc):
    sec = doc.sections[0]
    par = sec.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("SGMC | Definición funcional y mesa de trabajo | Concesión Transversal del Sisga S.A.S. | Página ")
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS
    r.font.name = "Segoe UI"
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    par._p.append(fld)


# --------------------------------------------------------------- contenido
DECISIONES = [
    dict(cod="D-01", tit="Coordenadas reales de los activos", bloque="A. Datos maestros",
         critica=True,
         enc="Los 34 activos del inventario tienen exactamente la misma coordenada: 4.728512, -74.114531. "
             "Ese punto está en Bogotá, no en el corredor.",
         imp="Toda la promesa de evidencia verificable del sistema descansa en el control GPS, y ese control compara la "
             "posición del técnico contra esa coordenada. Con el dato actual, un técnico parado frente a un poste SOS en "
             "Machetá no podrá cerrar nunca su mantenimiento, y en cambio cualquiera ubicado en ese punto de Bogotá "
             "validaría los 34 activos. La regla está bien programada; el dato la vuelve inútil.",
         prop="Levantar la coordenada de los 34 activos en un recorrido de corredor, capturando con el mismo celular que "
              "usará el técnico, y cargarlas al inventario. Aprovechar el recorrido para verificar que el código QR físico "
              "esté instalado y legible.",
         preg=["¿Quién hace el recorrido y en qué fecha?",
               "¿Existe ya un levantamiento topográfico o un inventario georreferenciado del que podamos tomar las "
               "coordenadas sin salir a campo?",
               "¿Los 34 activos cargados son todo el inventario, o falta registrar más?"],
         desb="El control GPS, el piloto de campo y los criterios de aceptación CA-01.3 y CA-01.4.", lineas=5),

    dict(cod="D-02", tit="Radio de tolerancia del control GPS", bloque="A. Datos maestros",
         enc="El radio definido es 1,0 km, heredado de la especificación original.",
         imp="Un kilómetro es mucho para un poste SOS en vía abierta, donde 100 metros bastarían para probar presencia. "
             "Y puede ser poco para un activo lineal como un tramo de fibra óptica, donde el técnico puede intervenir "
             "legítimamente a varios kilómetros del punto registrado. Un radio único para 18 tipos de activo muy distintos "
             "es una simplificación que conviene revisar.",
         prop="Radio diferenciado por tipo de activo: 200 m para activos puntuales (SOS, CCTV, paneles, básculas) y un "
              "tratamiento distinto para fibra óptica y activos lineales.",
         preg=["¿El kilómetro es una exigencia contractual o de interventoría, o es un valor elegido por el equipo?",
               "¿Acepta radio diferenciado por tipo de activo?"],
         desb="La configuración definitiva de la regla de validación.", lineas=4),

    dict(cod="D-03", tit="Qué significa la sede de un activo", bloque="A. Datos maestros",
         critica=True,
         enc="Los 11 usuarios están registrados en la sede 1, que es el CCO. Los 34 activos están repartidos entre las "
             "sedes 7 a 10, que son las unidades funcionales UF1 a UF4. La intersección es vacía.",
         imp="El filtro de seguridad descarga al celular del técnico solo los activos de su sede. Con los datos actuales, "
             "cada técnico descargaría cero activos y la aplicación quedaría vacía en sus manos el primer día de piloto. "
             "El problema de fondo es que la misma columna se usa con dos significados: para los usuarios es el sitio "
             "físico donde trabajan (CCO, peaje, báscula) y para los activos es el tramo del corredor donde están "
             "(UF1 a UF4).",
         prop="Separar los dos conceptos: la unidad funcional es un atributo del activo y la zona de trabajo es un "
              "atributo del usuario. Un técnico puede tener asignada más de una unidad funcional.",
         preg=["¿Un técnico atiende una unidad funcional completa, un tramo, o todo el corredor?",
               "¿Un técnico del CCO puede intervenir activos de cualquier unidad funcional?",
               "¿Qué debe ver un supervisor: su zona o todo el corredor?"],
         desb="El filtro de seguridad y, con él, la descarga de datos al celular. Sin esto el piloto no arranca.", lineas=5),

    dict(cod="D-04", tit="Qué hacer cuando el GPS falla", bloque="B. Operación en campo",
         enc="El sistema captura la precisión del GPS pero no define qué hacer cuando esa precisión es mala. El plan "
             "original ya señalaba el riesgo de falsos negativos y proponía un mecanismo de excepción supervisada, que "
             "nunca se especificó.",
         imp="En túnel, en cañón o bajo copa densa, el celular puede reportar una posición con error de cientos de metros "
             "o no fijar nada. Si el sistema bloquea sin salida, el técnico pierde el trabajo de una hora y el sistema "
             "pierde su confianza el primer día. Si el sistema deja pasar todo, el control no sirve para nada. Este es el "
             "punto donde el control se gana o se pierde.",
         prop="Permitir el cierre con excepción cuando la precisión reportada sea peor que un umbral, exigiendo al técnico "
              "un motivo escrito y una fotografía del activo, marcando el registro como cerrado con excepción y "
              "notificando al supervisor. La excepción no se oculta: queda visible en el reporte.",
         preg=["¿Acepta el cierre con excepción, o prefiere bloqueo estricto y que el técnico regrese?",
               "Si lo acepta, ¿quién autoriza: se marca solo y el supervisor revisa después, o requiere autorización "
               "previa por llamada al CCO?",
               "¿Qué porcentaje de excepciones sería aceptable antes de considerarse un problema?"],
         desb="La excepción E-02 del flujo de campo y buena parte de la aceptación del sistema por los técnicos.", lineas=5),

    dict(cod="D-05", tit="Interrupción a mitad de formulario", bloque="B. Operación en campo",
         enc="No hay definición de qué ocurre si el técnico cierra la aplicación, se le agota la batería o lo interrumpen "
             "a mitad del checklist.",
         imp="Una inspección de poste SOS son 15 preguntas más fotografías. Perder eso por una llamada entrante es la "
             "clase de fricción que hace que un técnico vuelva al papel.",
         prop="Guardado del borrador local, con la inspección visible como en curso hasta que se cierre formalmente.",
         preg=["¿Un mantenimiento puede quedar abierto de un día para otro, o debe cerrarse en la jornada?",
               "¿Un mantenimiento iniciado y no cerrado cuenta como incumplimiento?"],
         desb="La excepción E-03 y la definición del indicador de cumplimiento.", lineas=4),

    dict(cod="D-06", tit="Ciclo de vida de la orden de trabajo", bloque="B. Operación en campo",
         enc="Las órdenes registradas tienen estados Asignada, Cerrada y Suspendida, pero no hay regla que diga quién "
             "cambia cada estado ni bajo qué condición. Además, el único checklist existente apunta a una orden que no "
             "existe, señal de que hoy se puede registrar trabajo sin orden válida.",
         imp="Sin ciclo de vida definido no hay indicador de cumplimiento confiable: no se puede medir ejecutado contra "
             "programado si no está claro cuándo una orden cuenta como vencida, quién la suspende y si un trabajo puede "
             "existir sin orden.",
         prop="Estados: Programada, Asignada, En ejecución, En revisión, Cerrada, Suspendida y Vencida. El técnico mueve "
              "hasta En revisión; solo el supervisor cierra o suspende; el sistema marca Vencida por fecha. Ver figura 4.",
         preg=["¿Un técnico puede ejecutar un mantenimiento sin orden previa? Es el caso típico del correctivo detectado "
               "en ruta.",
               "¿Cuándo se considera vencida una orden: al día siguiente de la fecha programada, al cierre de semana, al "
               "cierre de mes?",
               "¿Quién puede suspender una orden y qué motivos son válidos?",
               "¿Se permite reasignar una orden a otro técnico, y queda traza de la reasignación?"],
         desb="El indicador de cumplimiento, el flujo de asignación y la excepción E-06.", lineas=5),

    dict(cod="D-07", tit="Trabajo incompleto, segunda visita y devoluciones", bloque="B. Operación en campo",
         enc="El modelo prevé marcar que se requiere una segunda visita con su motivo, y la aprobación del supervisor, "
             "pero ninguno de los dos flujos está definido.",
         imp="El caso real más común no es el mantenimiento perfecto: es el técnico que llega y no puede terminar por "
             "falta de repuesto, por lluvia o por acceso cerrado. Si el sistema no modela ese caso, el técnico va a "
             "forzar un cierre falso o no va a registrar nada.",
         prop="Cierre parcial con motivo tipificado, que genera automáticamente una orden de seguimiento asociada a la "
              "original. Y devolución del supervisor con observación, que reabre el mantenimiento al mismo técnico "
              "conservando la traza del rechazo.",
         preg=["¿Qué motivos de trabajo incompleto son válidos? Necesitamos la lista para el desplegable.",
               "¿La segunda visita es una orden nueva o la misma orden reabierta?",
               "Cuando el supervisor devuelve, ¿el técnico corrige el registro original o crea uno nuevo?",
               "¿El rechazo debe notificar por correo?"],
         desb="La excepción E-04, el flujo de aprobación y el cálculo real de cumplimiento.", lineas=5),

    dict(cod="D-08", tit="Activos no inventariados y correctivo desde campo", bloque="B. Operación en campo",
         enc="No hay ruta para un activo que el técnico encuentra en vía y que no está en el sistema, ni para una falla "
             "detectada fuera de programación.",
         imp="Sin esa ruta, los hallazgos de campo se pierden o se gestionan por WhatsApp, que es exactamente lo que el "
             "sistema viene a reemplazar.",
         prop="Permitir al técnico levantar un reporte de novedad con foto, coordenada y descripción, que llega al "
              "supervisor como solicitud de alta de activo o de orden correctiva, sin que el técnico pueda crear activos "
              "directamente.",
         preg=["¿Los técnicos deben poder reportar novedades de activos no inventariados?",
               "Un activo que queda fuera de servicio, ¿debe generar automáticamente una orden correctiva, o el "
               "correctivo se gestiona por fuera del sistema?"],
         desb="La excepción E-05 y el alcance de la alerta automática.", lineas=4),

    dict(cod="D-09", tit="Cuántos tipos de activo entran al primer sprint", bloque="C. Alcance de los formularios",
         critica=True,
         enc="Hay 18 tipos de activo y 18 formularios declarados, pero solo el de postes SOS tiene sus preguntas "
             "construidas: 15 preguntas. Los otros 17 están vacíos. Además, la columna que conecta cada tipo de activo "
             "con su formulario está vacía en los 18 tipos, de modo que hoy la aplicación no sabría qué checklist abrir "
             "ni siquiera para el SOS.",
         imp="Construir los 17 bancos restantes son del orden de 250 preguntas que alguien con criterio técnico debe "
             "redactar, con sus rangos, unidades y obligatoriedad. Es la tarea más grande que queda y es trabajo del "
             "equipo de la Concesión, no de configuración. Intentar construir los 18 a la vez es lo que hará que el "
             "proyecto se estanque otro mes.",
         prop="Arrancar con tres tipos que cubran la mayor cantidad de activos y de casuística (postes SOS, CCTV y "
              "paneles de mensaje variable), validar el ciclo completo en campo con ellos, e incorporar el resto en "
              "tandas quincenales según criticidad.",
         preg=["¿Cuáles tres tipos priorizamos? ¿Coincide con SOS, CCTV y paneles?",
               "¿Quién redacta y valida técnicamente las preguntas de cada tipo, y con qué disponibilidad?",
               "¿Existen formatos de inspección en papel vigentes que podamos transcribir en lugar de redactar desde "
               "cero? Sería la vía más rápida y la que mejor refleja la práctica real.",
               "¿Hay un requisito contractual o de interventoría sobre qué debe contener una inspección?"],
         desb="El alcance real del primer sprint y su fecha. Es la decisión que más mueve el cronograma.", lineas=6),

    dict(cod="D-10", tit="Evidencia fotográfica y firmas", bloque="C. Alcance de los formularios",
         enc="El requerimiento habla de hasta 6 fotografías. El modelo de datos las soporta de dos maneras a la vez, "
             "incompatibles entre sí: dos campos fijos dentro del registro de mantenimiento (imagen de inicio e imagen "
             "final) y una tabla separada sin límite. Lo mismo ocurre con las firmas. Las tablas separadas están vacías.",
         imp="Si no se decide, el técnico terminará firmando dos veces y adjuntando fotos en dos lugares distintos, o "
             "peor, la evidencia quedará repartida y los reportes no la encontrarán completa.",
         prop="Si se mantienen las 6 fotografías, la única vía viable es la tabla separada, y hay que retirar los campos "
              "fijos del registro de mantenimiento. Firma del técnico siempre; firma del supervisor solo si "
              "efectivamente firma en campo.",
         preg=["¿Cuántas fotografías exige realmente una inspección: un mínimo obligatorio y un máximo?",
               "¿Deben ser fotografías tipificadas (antes, después, novedad) o libres?",
               "¿El supervisor firma en campo junto al técnico, o su validación es la aprobación en el portal?",
               "¿La firma tiene valor contractual frente a interventoría, o es control interno?"],
         desb="El diseño definitivo del formulario y el cumplimiento verificable de los requerimientos de evidencia.",
         lineas=5),

    dict(cod="D-11", tit="Trazabilidad histórica de las respuestas", bloque="C. Alcance de los formularios",
         enc="El detalle de las inspecciones guarda el texto de la pregunta, no su identificador. Si alguien reformula "
             "una pregunta, los registros anteriores dejan de ser comparables con los nuevos.",
         imp="Si el sistema debe mostrar la evolución de un activo en el tiempo, o demostrar ante interventoría que se "
             "aplicó el mismo criterio durante un periodo, esta trazabilidad es indispensable. Si el sistema solo debe "
             "dejar constancia de cada visita por separado, no lo es.",
         prop="Guardar el identificador de la pregunta junto al texto, y versionar el formulario cuando cambie.",
         preg=["¿Necesita comparar la misma pregunta a lo largo del tiempo para un activo?",
               "¿El sistema debe poder reconstruir cómo era un formulario en una fecha pasada?"],
         desb="La comparabilidad histórica y los reportes de evolución.", lineas=4),

    dict(cod="D-12", tit="Qué reportes debe entregar el sistema", bloque="D. Reportes",
         critica=True,
         enc="Ningún reporte está definido. Se mencionan indicadores en documentos previos pero ninguno tiene fórmula, "
             "periodicidad ni destinatario.",
         imp="Este es el vacío más costoso del proyecto. Todo lo que el técnico captura en campo existe para producir "
             "algo, y ese algo nunca se especificó. Definir el reporte al final obliga casi siempre a volver atrás y "
             "capturar datos que no se pidieron. Definirlo ahora es lo que garantiza que el formulario pida lo correcto.",
         prop="Confirmar, descartar o completar la lista de reportes candidatos de la tabla siguiente, indicando para "
              "cada uno destinatario, periodicidad y formato.",
         preg=["¿Qué debe mostrar cada reporte y para tomar qué decisión?",
               "¿Quién lo recibe y con qué periodicidad: diario, semanal, mensual?",
               "¿En qué formato: pantalla, PDF por correo, Excel exportable?",
               "¿Alguno se entrega a un tercero, interventoría o ANI, con formato obligatorio?"],
         desb="El tablero y los entregables, y hacia atrás, la validación de que los formularios capturan lo necesario.",
         lineas=6, tabla_extra=True),

    dict(cod="D-13", tit="Definición de los indicadores", bloque="D. Reportes",
         enc="Se habla de cumplimiento, disponibilidad y tiempo de atención sin fórmula acordada.",
         imp="La disponibilidad de activos puede significar tres cosas distintas según se mida por tiempo, por cantidad "
             "o ponderada por criticidad, y cada una da un número diferente ante la misma realidad. Si ese número va a un "
             "informe de interventoría, la definición debe estar acordada antes y no después.",
         prop="Fijar por escrito la fórmula de cada indicador antes de construir el tablero.",
         preg=["Cumplimiento: ¿se calcula sobre órdenes cerradas en fecha, o basta con que se hayan ejecutado? ¿Una orden "
               "cerrada con excepción de GPS cuenta como cumplida?",
               "Disponibilidad: ¿por tiempo fuera de servicio, por cantidad de activos, o ponderada por criticidad?",
               "¿Existe una meta contractual de disponibilidad o de tiempo de atención frente a la ANI que el sistema "
               "deba reportar?",
               "¿Interventoría tendrá acceso al sistema, o solo recibe reportes exportados?"],
         desb="El tablero de indicadores y cualquier reporte a terceros.", lineas=5),

    dict(cod="D-14", tit="Usuarios, licenciamiento y gobierno del cambio", bloque="E. Gobierno",
         enc="Hay 11 usuarios registrados, dos de ellos inactivos. El presupuesto declarado en el plan original fue de "
             "100 USD mensuales, y la plataforma se cobra por usuario activo. No hay definición de quién puede modificar "
             "la aplicación en producción.",
         imp="El costo escala con el número de usuarios, y sin un responsable de cambios cualquiera puede alterar un "
             "formulario en producción y romper la comparabilidad de los datos.",
         prop="Definir el número de usuarios en régimen, el responsable funcional que autoriza cambios y la política de "
              "retención de evidencia.",
         preg=["¿Cuántos técnicos usarán el sistema en régimen, más allá de los 10 del piloto?",
               "¿El personal de interventoría o de la ANI tendrá acceso, aunque sea de consulta?",
               "¿Quién es el responsable funcional que autoriza cambios en producción?",
               "¿Cuánto tiempo debe conservarse la evidencia fotográfica y dónde se respalda?"],
         desb="El dimensionamiento de licencias y el procedimiento de control de cambios.", lineas=5),
]


def construir():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    est = doc.styles["Normal"]
    est.font.name = "Segoe UI"
    est.font.size = Pt(10.5)
    est.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")

    pie_de_pagina(doc)

    # ---------------------------------------------------------- PORTADA
    p(doc, "CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.", size=10, bold=True, color=AZUL, space_after=2)
    p(doc, "Sistema de Gestión de Mantenimiento en Campo", size=10, color=GRIS, space_after=28)

    p(doc, "Definición funcional", size=30, bold=True, color=TINTA, space_after=0)
    p(doc, "y mesa de trabajo", size=30, bold=True, color=AZUL, space_after=18)
    p(doc, "Documento de validación con el líder funcional", size=13, color=GRIS, space_after=30)

    tabla(doc, ["", ""], [
        ["Dirigido a", "Líder funcional y equipo de Operaciones / ITS"],
        ["Elaborado por", "Equipo técnico del SGMC"],
        ["Fecha", "6 de agosto de 2026"],
        ["Versión", "1.0"],
        ["Estado del proyecto", "Sprint 0. Definición funcional en validación"],
        ["Documento previo", "Plan de Implementación del SGMC sobre Google AppSheet, julio de 2026"],
        ["Qué se espera", "Respuestas a 14 decisiones y acta firmada"],
    ], anchos=[4.5, 12.0], size=10)

    p(doc, space_after=18)
    bloque(doc, "Este documento requiere su respuesta.",
           "No es un informe de avance ni una solicitud de aprobación. Es un cuestionario estructurado: catorce "
           "decisiones que hoy bloquean el proyecto y que solo el área funcional puede tomar. Cada decisión tiene una "
           "casilla para que escriba directamente sobre este archivo y lo devuelva.",
           "B0700C", AMBAR, "FDF3DE")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------- CÓMO USAR
    titulo(doc, "Cómo usar este documento", 1)
    p(doc, "Está organizado en tres partes que conviene recorrer en orden:")
    vinneta(doc, "Las secciones 1 a 4 describen cómo entendemos hoy que debe funcionar el sistema, expresado como flujos "
                 "de uso verificables. Léalas para confirmar o corregir.")
    vinneta(doc, "La sección 5 es el cuestionario. Cada punto sigue el mismo molde: qué encontramos, por qué importa, qué "
                 "proponemos y qué necesitamos que responda. Debajo de cada uno hay una casilla de respuesta.")
    vinneta(doc, "Las secciones 6 y 7 muestran cómo sus respuestas se convierten en trabajo y en fecha, y recogen el acta "
                 "de la sesión.")
    p(doc, space_after=4)
    p(doc, "Sugerencia de mesa: dos sesiones de dos horas. La primera recorre las secciones 1 a 4 y valida los flujos. La "
           "segunda resuelve el cuestionario decisión por decisión. Lo que no se decida queda con el supuesto que está "
           "declarado en cada punto, y ese supuesto se vuelve vinculante para la implementación.")

    # ---------------------------------------------------------- 1. CONTEXTO
    titulo(doc, "1. Por qué volvemos a la definición funcional", 1)
    p(doc, "El SGMC se planteó en julio de 2026 como un producto mínimo viable de 8 días sobre Google AppSheet, a partir "
           "de una especificación de requerimientos y un formulario de levantamiento diligenciado el 25 de julio. Se "
           "construyó, y hoy existe una aplicación publicada con un modelo de datos de 24 tablas, 34 activos catalogados "
           "y órdenes de trabajo registradas.")
    p(doc, "En el camino el alcance mutó. El modelo pasó de 17 a 24 tablas, aparecieron dos maneras distintas de definir "
           "los formularios de inspección, y las fotografías y firmas quedaron modeladas por duplicado. Ninguna de esas "
           "mutaciones pasó por una validación funcional.")
    p(doc, "Una auditoría del 6 de agosto de 2026 verificó ocho hallazgos bloqueantes leyendo directamente el archivo de "
           "producción. Al revisarlos uno a uno aparece un patrón que explica este documento:")
    bloque(doc, "",
           "La mayoría de los bloqueantes no son errores de programación. Son preguntas de negocio que nadie respondió. "
           "No se puede configurar el filtro de seguridad sin saber si la sede de un activo es la unidad funcional o el "
           "peaje. No se puede cerrar el requerimiento de evidencia fotográfica sin saber cuántas fotos exige realmente "
           "una inspección.",
           "1F5B99", AZUL, "E2ECF6")
    p(doc, "Por eso el paso siguiente no es seguir construyendo. Es definir, validar con usted, y recién entonces "
           "ejecutar.", space_before=6)

    # ---------------------------------------------------------- 2. HALLAZGOS
    titulo(doc, "2. Qué encontramos", 1)
    p(doc, "Verificado el 6 de agosto de 2026 leyendo el archivo maestro de producción. Lo que está construido y "
           "funcionando: el modelo de 24 tablas, los catálogos poblados con 34 activos y sus códigos QR, 18 tipos de "
           "activo, 10 sedes, 11 usuarios, 6 órdenes de trabajo, un checklist de inspección con su detalle, el banco de "
           "preguntas del formulario de postes SOS y la aplicación publicada con sus vistas móviles y web.")
    p(doc, "Lo que está abierto:")
    tabla(doc, ["#", "Hallazgo", "Decisión asociada"], [
        ["B-01", "Los 34 activos comparten una sola coordenada, situada en Bogotá y no en el corredor. El control GPS es "
                 "inoperante hasta levantar las coordenadas reales", "D-01"],
        ["B-02", "La columna que conecta cada tipo de activo con su formulario está vacía en los 18 tipos: la asignación "
                 "automática de checklist no tiene mapeo", "D-09"],
        ["B-03", "Todos los usuarios están en la sede 1 y todos los activos en las sedes 7 a 10. El filtro de seguridad "
                 "dejaría a cada técnico sin activos", "D-03"],
        ["B-04", "Solo 1 de 18 formularios tiene su banco de preguntas construido", "D-09"],
        ["B-05", "El único checklist existente referencia una orden de trabajo que no existe", "D-06"],
        ["B-06", "Fotografías, firmas y posición GPS están modelados dos veces: campos dentro del mantenimiento y tablas "
                 "separadas vacías", "D-10"],
        ["B-07", "La tabla de mantenimientos está vacía: ningún mantenimiento se ha ejecutado nunca de extremo a extremo",
         "Todas"],
        ["B-08", "El detalle de checklist guarda las preguntas como texto libre, sin trazabilidad al banco de preguntas",
         "D-11"],
    ], anchos=[1.6, 11.4, 3.5])

    p(doc, "El hallazgo B-01 merece verse en detalle, porque es el que invalida la promesa central del sistema.",
      space_before=6)
    figura(doc, "fig_03_coordenadas.png",
           "Figura 3. Los 34 activos tienen registrada la misma coordenada, ubicada en Bogotá y no en el corredor.")

    # ---------------------------------------------------------- 3. ACTORES
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "3. Actores y alcance", 1)
    p(doc, "Cuatro roles operan el sistema. Confirme que la descripción corresponde a la organización real y que no falta "
           "ningún actor.")
    figura(doc, "fig_01_actores.png", "Figura 1. Actores del sistema, ámbito de trabajo y restricción de datos.")

    titulo(doc, "Alcance propuesto para el primer sprint productivo", 2)
    p(doc, "Entra: mantenimiento preventivo programado, sobre los tipos de activo que se prioricen, ejecutado por técnico "
           "con evidencia fotográfica, firma y validación GPS, con asignación y aprobación por parte del supervisor, y un "
           "tablero básico de cumplimiento.")
    p(doc, "No entra por ahora, y se confirma en la mesa: mantenimiento correctivo iniciado desde el campo sin orden "
           "previa, generación automática de órdenes por frecuencia, integración con Power BI o mesas de ayuda, gestión "
           "de repuestos e inventario, y firma de interventoría.")
    p(doc, "Este corte es una propuesta. Las decisiones D-08, D-09 y D-12 lo redefinen.", italic=True, color=GRIS)

    # ---------------------------------------------------------- 4. FLUJOS
    titulo(doc, "4. Flujos funcionales", 1)
    p(doc, "Cada flujo se expresa con precondiciones, pasos, excepciones y criterios de aceptación verificables, de modo "
           "que sirva a la vez como definición funcional y como base de las pruebas de aceptación de usuario.")

    titulo(doc, "CU-01. El técnico ejecuta un mantenimiento en campo", 2, color=AZUL)
    p(doc, "Es el flujo que justifica el sistema. Los pasos marcados en ámbar dependen de una decisión pendiente.")
    figura(doc, "fig_02_flujo_tecnico.png",
           "Figura 2. Ciclo del técnico en campo y excepciones sin respuesta definida.")

    p(doc, "Criterios de aceptación", bold=True, space_before=6, space_after=4)
    tabla(doc, ["Código", "El sistema se considera conforme cuando"], [
        ["CA-01.1", "El registro se completa de principio a fin sin señal celular en ningún momento"],
        ["CA-01.2", "Al recuperar señal, el registro aparece en el sistema sin acción del técnico"],
        ["CA-01.3", "Estando a más de la distancia definida del activo, el sistema impide guardar"],
        ["CA-01.4", "Estando junto al activo, el sistema permite guardar"],
        ["CA-01.5", "El checklist que abre corresponde al tipo del activo escaneado"],
        ["CA-01.6", "Las fotografías quedan asociadas al mantenimiento y comprimidas"],
        ["CA-01.7", "Un mantenimiento sin firma no puede cerrarse"],
    ], anchos=[2.6, 13.9])
    bloque(doc, "Advertencia sobre el estado actual.",
           "Este flujo nunca se ha ejecutado completo: la tabla de mantenimientos está vacía. Y como los 34 activos "
           "comparten una coordenada en Bogotá, hoy fallan a la vez CA-01.3 y CA-01.4.",
           "A82D2D", ROJO, "FBE9E9")

    titulo(doc, "CU-02. El supervisor programa y asigna una orden de trabajo", 2, color=AZUL)
    p(doc, "El supervisor crea la orden en el portal seleccionando activo, técnico, fecha programada y tipo de trabajo. "
           "El sistema asigna el número, un bot notifica por correo al técnico y la orden aparece en su celular al "
           "sincronizar.")
    p(doc, "Sin resolver: los estados que hoy existen en los datos son Asignada, Cerrada y Suspendida, pero no hay "
           "definición de quién puede pasar de uno a otro ni bajo qué condición. Es la decisión D-06.", color=AMBAR)
    figura(doc, "fig_04_ciclo_ot.png", "Figura 4. Ciclo de vida propuesto para la orden de trabajo.")

    titulo(doc, "CU-03. El supervisor revisa y aprueba", 2, color=AZUL)
    p(doc, "Al sincronizar el técnico, el mantenimiento aparece en el portal. El supervisor revisa checklist, "
           "fotografías, firma y coordenada de cierre, y aprueba o devuelve con observación. Al aprobar, la orden se "
           "cierra y el indicador de cumplimiento se actualiza.")
    p(doc, "Sin resolver: no existe el flujo de rechazo. Qué pasa con un mantenimiento devuelto, si el técnico lo corrige "
           "sobre el mismo registro o crea uno nuevo, y si el rechazo notifica. Es la decisión D-07.", color=AMBAR)

    titulo(doc, "CU-04. Alerta automática por activo fuera de servicio", 2, color=AZUL)
    p(doc, "Cuando el técnico cierra un mantenimiento con estado final Fuera de servicio, al sincronizar se dispara un "
           "bot que genera un informe en PDF con los datos de la falla, la ubicación y las fotografías, y lo envía por "
           "correo prioritario al CCO y al supervisor de la zona.")
    p(doc, "Sin resolver: a quién exactamente se notifica, si hay escalamiento cuando nadie responde, y si un activo "
           "fuera de servicio debe generar automáticamente una orden correctiva. Es la decisión D-08.", color=AMBAR)

    titulo(doc, "CU-05. El administrador mantiene el inventario y los formularios", 2, color=AZUL)
    p(doc, "Da de alta activos con su código, tipo, sede, punto de referencia, calzada, sentido, coordenada real, "
           "frecuencia y código QR. Define o ajusta el formulario de inspección de cada tipo de activo y publica el "
           "cambio, que los técnicos reciben en la siguiente sincronización.")
    p(doc, "Sin resolver: modificar una pregunta hoy rompe la comparación histórica, porque el detalle guarda el texto y "
           "no el identificador de la pregunta. Es la decisión D-11.", color=AMBAR)

    titulo(doc, "CU-06. Reportes y tablero", 2, color=ROJO)
    p(doc, "Este caso de uso no está definido, y es el vacío más grande del proyecto: el sistema captura datos desde "
           "julio pero nadie ha especificado qué debe producir con ellos. Lo que se menciona en documentos previos "
           "(cumplimiento, tiempo de atención, disponibilidad por zona, mapa de activos fuera de servicio) no tiene "
           "fórmula, periodicidad, destinatario ni formato acordado. Son las decisiones D-12 y D-13.")

    # ---------------------------------------------------------- 5. DECISIONES
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "5. Las catorce decisiones", 1)
    p(doc, "Agrupadas en cinco bloques. Cada una bloquea trabajo concreto. Las marcadas como ruta crítica son las que "
           "fijan el cronograma del proyecto.")

    bloque_actual = None
    for d in DECISIONES:
        if d["bloque"] != bloque_actual:
            bloque_actual = d["bloque"]
            titulo(doc, "Bloque " + bloque_actual, 2, color=GRIS, size=11)

        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(14)
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.keep_with_next = True
        r = par.add_run(d["cod"] + ".  ")
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = AZUL; r.font.name = "Segoe UI"
        r = par.add_run(d["tit"])
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = TINTA; r.font.name = "Segoe UI"
        if d.get("critica"):
            r = par.add_run("     RUTA CRÍTICA")
            r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = ROJO; r.font.name = "Segoe UI"

        p(doc, "Qué encontramos", size=9.5, bold=True, color=GRIS, space_before=6, space_after=2)
        p(doc, d["enc"], space_after=6)
        p(doc, "Por qué importa", size=9.5, bold=True, color=GRIS, space_after=2)
        p(doc, d["imp"], space_after=6)
        p(doc, "Qué proponemos", size=9.5, bold=True, color=GRIS, space_after=2)
        p(doc, d["prop"], space_after=6)

        if d.get("tabla_extra"):
            p(doc, "Reportes candidatos. Confirme, descarte o complete:", size=9.5, bold=True, color=GRIS, space_after=4)
            tabla(doc, ["Reporte", "Para qué sirve", "Destinatario", "Periodicidad", "¿Va?"], [
                ["Cumplimiento del plan", "Ejecutado contra programado, por zona y por técnico", "Supervisor, Dirección", "Mensual", ""],
                ["Activos fuera de servicio", "Qué está caído, desde cuándo y quién lo atiende", "CCO", "Diario", ""],
                ["Hoja de vida del activo", "Historial completo de intervenciones de un equipo", "Supervisor, Interventoría", "A demanda", ""],
                ["Certificado de mantenimiento", "Constancia de una intervención con evidencia y firma", "Interventoría", "Por intervención", ""],
                ["Productividad del técnico", "Mantenimientos ejecutados y tiempo promedio", "Supervisor", "Semanal", ""],
                ["Excepciones de GPS", "Cierres fuera de rango o de baja precisión", "Supervisor", "Semanal", ""],
            ], anchos=[3.6, 5.4, 3.4, 2.4, 1.7], size=9)

        p(doc, "Necesitamos que responda", size=9.5, bold=True, color=AMBAR, space_after=3)
        for q in d["preg"]:
            vinneta(doc, q, size=10.5)

        casilla_respuesta(doc, alto_lineas=d.get("lineas", 5))

        p(doc, "Desbloquea: " + d["desb"], size=9, italic=True, color=GRIS, space_after=10)

    # ---------------------------------------------------------- 6. ROADMAP
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "6. De sus respuestas al cronograma", 1)
    p(doc, "El cronograma no está escrito porque depende de esta mesa. Este es el mapa de dependencias: qué habilita "
           "cada decisión y en qué orden se ejecuta el trabajo.")
    figura(doc, "fig_05_ruta_critica.png",
           "Figura 5. Secuencia de ejecución y ruta crítica del proyecto.")

    p(doc, "La ruta crítica son D-01 y D-09. Levantar las coordenadas de los 34 activos y redactar los bancos de "
           "preguntas son trabajo del equipo de la Concesión, no de configuración, y ambas se miden en semanas. El "
           "cronograma completo lo fijan esas dos y no las demás.", bold=True, space_before=4)

    titulo(doc, "Orden de ejecución una vez cerrada la mesa", 2)
    for i, (t, s) in enumerate([
        ("Definición", "Esta mesa y sus respuestas, con acta firmada."),
        ("Datos", "Coordenadas reales, mapeo de formulario por tipo de activo, realineación de sedes y construcción de "
                  "los bancos de preguntas priorizados."),
        ("Configuración", "Reglas de validación, filtro de seguridad, formularios, bots de notificación y reportes."),
        ("Prueba controlada", "Un mantenimiento completo de extremo a extremo ejecutado por una persona, primero con "
                              "señal y luego en modo avión. El criterio de cierre es que existan registros reales en la "
                              "base, verificados leyendo el archivo."),
        ("Piloto", "Los 10 celulares en vía, solo después de que la prueba controlada haya funcionado."),
        ("Producción y evolución", "Resto de tipos de activo, integraciones y respaldo automático."),
    ], 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(4)
        r = par.add_run(f"{i}. {t}. ")
        r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = AZUL; r.font.name = "Segoe UI"
        r = par.add_run(s)
        r.font.size = Pt(10.5); r.font.color.rgb = TINTA; r.font.name = "Segoe UI"

    # ---------------------------------------------------------- 7. ACTA
    titulo(doc, "7. Acta de la mesa de trabajo", 1)
    p(doc, "Para diligenciar al cierre de la sesión.")
    tabla(doc, ["Campo", "Registro"], [
        ["Fecha de la sesión", ""],
        ["Participantes", ""],
        ["Decisiones cerradas", ""],
        ["Decisiones aplazadas y hasta cuándo", ""],
        ["Supuestos que quedan vinculantes", ""],
        ["Responsable de D-01 (coordenadas) y fecha comprometida", ""],
        ["Responsable de D-09 (bancos de preguntas) y fecha comprometida", ""],
        ["Próxima sesión", ""],
    ], anchos=[6.5, 10.0], size=10)

    p(doc, space_after=24)
    t = doc.add_table(rows=1, cols=2)
    for i, et in enumerate(["Líder funcional", "Por el equipo técnico"]):
        c = t.cell(0, i)
        c.paragraphs[0].text = ""
        pp = c.paragraphs[0]
        r = pp.add_run("\n\n____________________________________")
        r.font.size = Pt(10.5); r.font.name = "Segoe UI"; r.font.color.rgb = TINTA
        pp2 = c.add_paragraph()
        r = pp2.add_run(et + "\nNombre y fecha")
        r.font.size = Pt(9); r.font.color.rgb = GRIS; r.font.name = "Segoe UI"

    p(doc, space_after=14)
    p(doc, "Las decisiones registradas en esta acta son la base del cronograma de implementación y del alcance del "
           "sprint.", size=9, italic=True, color=GRIS)

    doc.save(SALIDA)
    print("Documento generado:", SALIDA)


if __name__ == "__main__":
    construir()
