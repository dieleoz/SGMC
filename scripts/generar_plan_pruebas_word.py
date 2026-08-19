# -*- coding: utf-8 -*-
"""
Generador del Plan de Pruebas Periciales y Carga de Datos Técnicos en formato Word (.docx)
para la Concesión Transversal del Sisga S.A.S. (SGMC v2).
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generar_plan_pruebas_word(output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Estilos Base
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(15, 23, 42)

    # Encabezado Oficial
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    run_concesion = title_p.add_run("CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.")
    run_concesion.bold = True
    run_concesion.font.size = Pt(13)
    run_concesion.font.color.rgb = RGBColor(4, 120, 87)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run("SISTEMA DE GESTIÓN DE MANTENIMIENTO EN CAMPO (SGMC v2)\nPLAN DE PRUEBAS PERICIALES Y CARGA DE DATOS TÉCNICOS (27 FICHAS)")
    run_sub.bold = True
    run_sub.font.size = Pt(15)
    run_sub.font.color.rgb = RGBColor(15, 23, 42)

    # Tabla de Metadatos
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    headers_meta = [
        ("Código:", "PLAN-PRUEBAS-V2"),
        ("Corredor:", "137 km (Sisga - Aguaclara)"),
        ("Fichas Evaluadas:", "27 Subsistemas / 333 Ítems"),
        ("Emisión:", "Agosto 2026")
    ]
    
    for i, (label, val) in enumerate(headers_meta):
        cell = meta_table.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label} ")
        r1.bold = True
        r1.font.size = Pt(8.5)
        r2 = p.add_run(val)
        r2.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 1. Objetivo
    doc.add_heading("1. Objetivo del Plan de Pruebas", level=1)
    doc.add_paragraph(
        "Establecer el protocolo técnico y pericial para la verificación exhaustiva, validación de tolerancias instrumentales y carga controlada de datos de prueba sobre las 27 fichas de inspección técnica asociadas a los 368 activos de la Concesión Transversal del Sisga.\n\n"
        "El protocolo asegura que cada formulario capture con precisión física y matemática los parámetros exigidos en el Apéndice Técnico 1 de la ANI (voltajes, frecuencias, potencias, atenuaciones ópticas, presiones y estados operativos)."
    ).paragraph_format.space_after = Pt(10)

    # 2. Matriz de Parámetros por Ficha
    doc.add_heading("2. Matriz de Parámetros Instrumentales por Subsistema (27 Fichas)", level=1)
    
    fichas_table = doc.add_table(rows=1, cols=5)
    fichas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fichas_table.autofit = False

    cols_w = [Inches(1.0), Inches(1.6), Inches(1.8), Inches(1.1), Inches(1.5)]
    headers_fichas = ["Código", "Subsistema", "Parámetros Clave a Medir", "Tolerancia", "Criterio Interventoría"]
    
    for i, h_text in enumerate(headers_fichas):
        cell = fichas_table.rows[0].cells[i]
        cell.width = cols_w[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    fichas_data = [
        ("FRM_SOS", "Poste SOS (Auxilio)", "• Tensión batería DC\n• Tensión panel solar\n• Audio Full-Duplex CCO\n• Hermeticidad gabinete", "12.0V - 14.5V\n18.0V - 22.0V\nClaridad >95%\nGrado IP66", "Batería >=12.4V y audio nítido con el CCO."),
        ("FRM_CCTV", "Cámaras CCTV", "• Tensión PoE+ (Inyector)\n• Tasa de cuadros (FPS)\n• Movimiento PTZ 360°\n• Iluminador infrarrojo", "48.0V - 54.0V\n>=25 FPS\nPreset <1.5s\nAlcance >=150m", "Flujo RTSP continuo al CCO sin pérdidas."),
        ("FRM_PMVF", "Panel Mensaje Fijo", "• Alimentación principal AC\n• Voltaje banco DC\n• Conmutación píxeles LED\n• Protocolo NTCIP", "110V - 220V AC\n24.0V / 48.0V\n100% LEDs OK\nRespuesta <500ms", "Mensaje legible a 250m con fotocelda."),
        ("FRM_PMVM", "Panel Mensaje Móvil", "• Carga batería remolque\n• Elevación hidráulica\n• Enlace 4G/LTE CCO", "12.0V - 24.0V\nPresión normal\nLatencia <200ms", "Elevación y orientación en bermas."),
        ("FRM_GENE", "Grupos Electrógenos", "• Nivel diésel en tanque\n• Tensión arranque DC\n• Frecuencia de salida\n• Tensión trifásica\n• Tiempo conmutación ATS", ">=75% tanque\n24.0V - 27.5V\n60.0Hz ±0.5Hz\n208V / 220V AC\n<10 segundos", "Transferencia ATS inmediata ante corte de red."),
        ("FRM_UPS", "Sistemas UPS", "• Tensión por celda\n• Tensión banco total\n• Autonomía en descarga\n• Temperatura operación", "2.25V - 2.30V/c\n120V / 240V DC\n>=120 minutos\n20°C - 25°C", "Sin caída en transición a modo batería."),
        ("FRM_SUBE", "Subestaciones Eléctr.", "• Nivel aislamiento (MΩ)\n• Puesta a tierra\n• Aceite transformador", ">=1000 MΩ\n<=5.0 Ω\nNivel y rigidez", "Tierra <=5.0 Ω y termografía conforme."),
        ("FRM_FO", "Fibra Óptica (137km)", "• Atenuación por km\n• Pérdida por empalme\n• Reflectancia conectores", "<=0.25 dB/km\n<=0.05 dB/fusión\n<=-50 dB", "Margen óptico para anillo ERPS."),
        ("FRM_SWIT", "Switches de Campo", "• Enlace óptico SFP\n• Tasa de error tramas\n• Temperatura chasis", "-15 a -8 dBm\n0 errores CRC\n-20°C a +70°C", "Convergencia de anillo ITU-T <50 ms."),
        ("FRM_SWL3", "Switches Core Capa 3", "• Uso CPU / Memoria\n• Rutas BGP/OSPF\n• Fuentes redundantes", "CPU <35%\n100% adyacencias\nDual AC activa", "Cero interrupción de enrutamiento."),
        ("FRM_ROUT", "Enrutadores Borde", "• Túneles VPN IPsec\n• Throughput tráfico\n• Firewall / Seguridad", "AES-256\n>=1 Gbps\n0 accesos ilícitos", "Telemetría segura a Centro de Control."),
        ("FRM_SERV", "Servidores CCO", "• Estado arreglos RAID\n• Ocupación de disco\n• Servicios SGMC / BD", "RAID 10 OK\n<70% ocupación\nUptime >=99.9%", "Integridad de réplicas en PostgreSQL."),
        ("FRM_NAS", "Almacenamiento NAS", "• Retención video CCTV\n• Espacio libre en disco\n• Estado discos SMART", ">=30 días 1080p\n>=20 TB libres\nSMART OK", "Grabación continua de 30 días pliego ANI."),
        ("FRM_VW", "Video Wall CCO", "• Calibración cromática\n• Controladora matriz\n• Conmutación pantallas", "Sin píxeles OFF\nMatriz 4x2 activa\n<1 segundo", "Despliegue simultáneo de cámaras y alarmas."),
        ("FRM_ETD", "Estación Tráfico ETD", "• Conteo volumétrico\n• Clasificación por ejes\n• Sensor piezocable", "Precisión >=98%\nError <2%\nInductancia OK", "Calibración anual cotejada con aforo."),
        ("FRM_BASC", "Báscula Pesaje Est.", "• Error pesaje bruto\n• Celdas de carga\n• Plataforma estructural", "±0.5% error máx\nSimetría 4 celdas\nSin corrosión", "Certificado de pesas patrón vigente (SIC)."),
        ("FRM_BASD", "Báscula Dinámica WIM", "• Sensor piezo cuarzo\n• Error peso por eje\n• Detección sobrepeso", "Error <5% a 60km/h\nLectura continua\nTransmisión <1s", "Alerta automática en tiempo real al peaje."),
        ("FRM_OCR", "Cámaras OCR Placas", "• Tasa lectura placas\n• Iluminador infrarrojo\n• Integración RUNT", ">=96% placas\nSincronía obturador\nRespuesta <800ms", "Identificación y vinculación al pesaje."),
        ("FRM_PJC", "Peaje: Carril", "• Barrera electromecánica\n• Semáforo de carril\n• Sensor de altura", "Apertura <1.2s\nLEDs 100% OK\nDetección precisa", "Cero atascamiento en cabinas de peaje."),
        ("FRM_PJE", "Peaje: Telepeaje TAG", "• Antena RFID 915 MHz\n• Protocolo Colpass\n• Tasa lectura Tag", "Potencia nominal\nInteroperable\n>=99.5% lectura", "Cobro electrónico sin detención (Colpass)."),
        ("FRM_PSEG", "Paso Seguro Peatonal", "• Pulsador peatonal\n• Balizas luminosas\n• Sonorizador invidentes", "Contacto seco\n>=60 destellos/min\n85 dB a 1 metro", "Activación inmediata de advertencia vial."),
        ("FRM_SGE", "Gestión Energía Peaje", "• Factor de potencia\n• Banco condensadores\n• Supresores DPS", "FP >=0.95\nAutomático pasos\nVaristores intactos", "Cero penalización por energía reactiva."),
        ("FRM_SGM", "Gestión Ambiental", "• Sensor visibilidad\n• Pluviómetro lluvia\n• Velocidad de viento", "0 - 2000m\nResolución 0.2mm\n0 - 150 km/h", "Alertas meteorológicas en PMV por niebla."),
        ("FRM_SSA", "Señalización Acústica", "• Sirenas de túnel\n• Megafonía pública\n• Nivel presión sonora", "110 dB a 3m\nSTI >0.6\nClase D", "Evacuación audible en túneles del Sisga."),
        ("FRM_FIRE", "Extinción y Fuego", "• Sensor térmico lineal\n• Pulsadores emergencia\n• Gabinetes mangueras", "68°C / 88°C\nLazo direccionable\nPresión >=6 bar", "Activación de extractores en túnel."),
        ("FRM_PORT", "Portátiles Técnicos", "• Batería diagnóstico\n• Software de campo\n• Certificados VPN", ">=4 horas\nHerramientas ITS\nVigentes", "Terminal de configuración pericial."),
        ("FRM_IMPR", "Impresoras Térmicas", "• Cabezal térmico\n• Velocidad impresión\n• Sensor fin papel", "Corte <0.5s\n>=200 mm/s\nAviso preventivo", "Emisión continua de tiquetes de peaje.")
    ]

    for row_idx, r_data in enumerate(fichas_data):
        row = fichas_table.add_row()
        bg = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        for i, text in enumerate(r_data):
            cell = row.cells[i]
            cell.width = cols_w[i]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 60, 60)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            r.font.size = Pt(7.5)
            if i == 0:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. Protocolo de Ejecución
    doc.add_heading("3. Protocolo de Ejecución de Pruebas de Campo", level=1)
    doc.add_paragraph(
        "1. Calibración Instrumental: Uso de multímetro, telurómetro y OTDR calibrados con certificado vigente.\n"
        "2. Ejecución Offline en PWA (/tecnico): Captura satelital GPS con radio de tolerancia (50m/100m) o justificación de excepción en túneles. Registro de mediciones numéricas exactas, 2 fotos WebP georreferenciadas y firma digital manuscrita.\n"
        "3. Auditoría y Certificación (/supervisor): Contraste de geofencing en mapa PostGIS, verificación de fotos en Supabase Storage, aprobación y descarga de la Ficha Técnica Pericial en PDF."
    ).paragraph_format.space_after = Pt(10)

    # 4. Carga de Datos
    doc.add_heading("4. Protocolo de Carga de Datos Técnicos de Prueba", level=1)
    doc.add_paragraph(
        "Para ejecutar la batería y poblar las 27 fichas de prueba en lote con datos periciales exactos, ejecute en terminal:\n\n"
        "    python scripts/cargar_datos_pruebas_fichas.py\n\n"
        "El script asienta 27 órdenes cerradas con sus 333 respuestas en CHD_ChecklistDetalle, 54 fotos WebP georreferenciadas y 27 firmas digitales."
    ).paragraph_format.space_after = Pt(16)

    # 5. Firmas
    doc.add_heading("5. Certificación de Aprobación del Plan de Pruebas", level=1)
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False

    sig_w = [Inches(3.2), Inches(3.2)]
    sig_content = [
        ("POR LA CONCESIÓN TRANSVERSAL DEL SISGA:\n\n\n\n_____________________________________________\nIng. Diego Zúñiga\nCoordinador ITS / Director Técnico SGMC\nConcesión Transversal del Sisga S.A.S."),
        ("POR EL CONSORCIO DE INTERVENTORÍA:\n\n\n\n_____________________________________________\nIngeniero Especialista ITS / Auditor Contractual\nConsorcio Interventoría Sisga\nAgencia Nacional de Infraestructura (ANI)")
    ]

    for i, text in enumerate(sig_content):
        cell = sig_table.rows[0].cells[i]
        cell.width = sig_w[i]
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, 120, 120, 120, 120)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(text)
        r.font.size = Pt(8.5)

    doc.save(output_path)
    print(f"[OK] Plan de Pruebas Word generado exitosamente en: {output_path}")

if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Manuales", "PLAN_DE_PRUEBAS_Y_CARGA_DATOS_SGMC_V2.docx")
    generar_plan_pruebas_word(out)
