# -*- coding: utf-8 -*-
"""
Generador del Documento Oficial de Entrega:
Plan de Pruebas Funcionales (UAT), Validación por Rol y Certificación de Fichas Técnicas
en formato Microsoft Word (.docx) para la Concesión Transversal del Sisga S.A.S. (SGMC v2).
"""

import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=70, bottom=70, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generar_documento_word(output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(15, 23, 42)

    # -------------------------------------------------------------
    # PORTADA / ENCABEZADO
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    r_conc = title_p.add_run("CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.")
    r_conc.bold = True
    r_conc.font.size = Pt(13)
    r_conc.font.color.rgb = RGBColor(4, 120, 87)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("SISTEMA DE GESTIÓN DE MANTENIMIENTO EN CAMPO (SGMC v2)\nPLAN DE PRUEBAS DE ACEPTACIÓN FUNCIONAL POR ROL (UAT) Y VALIDACIÓN DE FICHAS TÉCNICAS")
    r_sub.bold = True
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(15, 23, 42)

    # Tabla de Metadatos
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    headers_meta = [
        ("Documento:", "PLN-UAT-SGMC-V2"),
        ("Corredor:", "137 km (Sisga - Aguaclara)"),
        ("Alcance:", "4 Roles / 27 Fichas"),
        ("Emisión:", "Agosto 2026")
    ]
    
    for i, (label, val) in enumerate(headers_meta):
        cell = meta_table.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 60, 60, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label} ")
        r1.bold = True
        r1.font.size = Pt(8)
        r2 = p.add_run(val)
        r2.font.size = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 1. OBJETIVO Y METODOLOGÍA
    # -------------------------------------------------------------
    doc.add_heading("1. Objetivo y Metodología de Validación Funcional (UAT)", level=1)
    doc.add_paragraph(
        "El presente documento define los casos de prueba formales de Aceptación de Usuario (UAT) para que cada rol funcional de la Concesión Transversal del Sisga e Interventoría valide de extremo a extremo la operatividad del SGMC v2 (https://sisga-2.vercel.app/).\n\n"
        "La metodología evalúa el ciclo completo de mantenimiento: ejecución offline en vía, supervisión geográfica pericial, programación preventiva masiva, control de disponibilidad contractual ante la ANI y la captura técnica de los 27 tipos de fichas de inspección."
    ).paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # 2. CASOS DE PRUEBA POR ROL FUNCIONAL
    # -------------------------------------------------------------
    doc.add_heading("2. Casos de Prueba Funcionales por Rol de Usuario", level=1)

    # Tabla de Casos de Prueba
    uat_table = doc.add_table(rows=1, cols=5)
    uat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    uat_table.autofit = False

    cols_w = [Inches(0.9), Inches(1.2), Inches(2.3), Inches(1.8), Inches(0.8)]
    headers_uat = ["Caso ID", "Rol Asignado", "Acción y Procedimiento", "Resultado Esperado", "Estado"]
    
    for i, h_text in enumerate(headers_uat):
        cell = uat_table.rows[0].cells[i]
        cell.width = cols_w[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(255, 255, 255)

    casos_data = [
        # TÉCNICO
        ("CP-TEC-01", "ROL-03 Técnico\n(Iván Salcedo)", "Abrir /tecnico en móvil y presionar 'Instalar App'.", "App se instala en inicio y opera a pantalla completa con caché offline.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-TEC-02", "ROL-03 Técnico\n(Iván Salcedo)", "Poner celular en Modo Avión. Abrir orden OT asignada.", "Carga formulario y preguntas sin conexión (Modo Offline Activo).", "[  ] Conforme\n[  ] No Conf."),
        ("CP-TEC-03", "ROL-03 Técnico\n(Iván Salcedo)", "Presionar 'Capturar Coordenadas GPS' junto al activo.", "Valida presencia satelital dentro del radio del equipo (ej. <=50m).", "[  ] Conforme\n[  ] No Conf."),
        ("CP-TEC-04", "ROL-03 Técnico\n(Luis Gacha)", "Activar 'Cierre con Excepción' en túnel o bajo techo.", "Permite cierre justificado dejando coordenadas en NULL (sin inventar).", "[  ] Conforme\n[  ] No Conf."),
        ("CP-TEC-05", "ROL-03 Técnico\n(Iván Salcedo)", "Diligenciar mediciones (ej. Batería SOS: 12.8V, GENE: 60Hz), 2 fotos WebP y firma.", "Guarda en IndexedDB, bloquea doble clic y muestra badge '1 pendiente'.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-TEC-06", "ROL-03 Técnico\n(Iván Salcedo)", "Restaurar red (desactivar Modo Avión).", "SyncEngine sincroniza solo en 3s y pasa a 'Cola al día (0 pendientes)'.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-TEC-07", "ROL-03 Técnico\n(Luis Gacha)", "En /novedades, reportar daño en vía con foto y GPS.", "Guarda en NOV_Novedades y crea OT Correctiva automática en Supabase.", "[  ] Conforme\n[  ] No Conf."),
        
        # SUPERVISOR
        ("CP-SUP-01", "ROL-02 Supervisor\n(Fernand Bolívar)", "Entrar a /supervisor y filtrar por UF1 y estado 'En revisión'.", "Muestra lista de órdenes enviadas por los técnicos con badge GPS.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-SUP-02", "ROL-02 Supervisor\n(Fernand Bolívar)", "Presionar 'Auditar Evidencias' en una orden cerrada.", "Despliega mapa PostGIS, fotos en S3, firma y checklist contestado.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-SUP-03", "ROL-02 Supervisor\n(Fernand Bolívar)", "Presionar 'Aprobar y Certificar'.", "Orden pasa a 'Cerrada' y se incorpora al cálculo de disponibilidad.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-SUP-04", "ROL-02 Supervisor\n(Fernand Bolívar)", "Presionar '📄 Descargar Ficha PDF (Interventoría)'.", "Genera Ficha Técnica Pericial con membrete oficial y fotos WebP.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-SUP-05", "ROL-02 Supervisor\n(Fernand Bolívar)", "En /planes, seleccionar mes y presionar 'Generar OTs del Mes'.", "Genera lote masivo de OTs preventivas en PLA_PlanMantenimiento.", "[  ] Conforme\n[  ] No Conf."),

        # DIRECTOR TÉCNICO / CCO
        ("CP-DIR-01", "ROL-01 Director\n(Diego Zúñiga)", "Ingresar al Centro de Control (/).", "Visualiza estado general del corredor (137 km) y avance de OTs.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-DIR-02", "ROL-01 Director\n(Diego Zúñiga)", "Ingresar a /activos y buscar por PK o código.", "Lista los 368 activos georreferenciados con formulario y tolerancias.", "[  ] Conforme\n[  ] No Conf."),

        # INTERVENTORÍA / ANI
        ("CP-INT-01", "ROL-04 Interventor\n(Consorcio Sisga)", "Ingresar a /reportes y consultar Disponibilidad (Di).", "Calcula fórmula contractual en vivo y muestra semáforo (>=98.5%).", "[  ] Conforme\n[  ] No Conf."),
        ("CP-INT-02", "ROL-04 Interventor\n(Consorcio Sisga)", "Revisar el Parte Diario de Operaciones del CCO.", "Consolida OTs programadas, ejecutadas, excepciones y novedades.", "[  ] Conforme\n[  ] No Conf."),
        ("CP-INT-03", "ROL-04 Interventor\n(Consorcio Sisga)", "Presionar 'Descargar Informe PDF (ANI)'.", "Emite Informe Oficial mensual con firmas listo para radicación ANI.", "[  ] Conforme\n[  ] No Conf.")
    ]

    for row_idx, r_data in enumerate(casos_data):
        row = uat_table.add_row()
        bg = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        for i, text in enumerate(r_data):
            cell = row.cells[i]
            cell.width = cols_w[i]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 50, 50, 50, 50)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            r.font.size = Pt(7.5)
            if i in [0, 4]:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # 3. MATRIZ INSTRUMENTAL DE LAS 27 FICHAS
    # -------------------------------------------------------------
    doc.add_heading("3. Matriz de Validación Instrumental de las 27 Fichas Técnicas", level=1)
    doc.add_paragraph(
        "A continuación se presenta la matriz de parámetros que deben medirse y validarse en campo para cada uno de los 27 subsistemas instalados a lo largo de los 137 km:"
    ).paragraph_format.space_after = Pt(4)

    fichas_table = doc.add_table(rows=1, cols=4)
    fichas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fichas_table.autofit = False

    cols_f_w = [Inches(1.0), Inches(1.8), Inches(2.4), Inches(1.8)]
    headers_f = ["Ficha ID", "Subsistema", "Parámetros Técnicos a Medir", "Tolerancia / Criterio ANI"]
    
    for i, h_text in enumerate(headers_f):
        cell = fichas_table.rows[0].cells[i]
        cell.width = cols_f_w[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(255, 255, 255)

    fichas_resumen = [
        ("FRM_SOS", "Postes SOS (Poste Auxilio)", "Tensión batería DC (12.8V), panel solar (19.5V), audio CCO.", "Batería >=12.4V y comunicación nítida con el CCO."),
        ("FRM_CCTV", "Cámaras CCTV", "Tensión PoE+ (52V), tasa 25 FPS 1080p, movimiento PTZ.", "Flujo RTSP continuo sin pérdida de paquetes."),
        ("FRM_PMVF", "Panel Mensaje Fijo", "Alimentación 220V AC, banco 24V DC, píxeles LED, NTCIP.", "Legible a 250m con fotocelda adaptativa."),
        ("FRM_PMVM", "Panel Mensaje Móvil", "Carga remolque 24V DC, elevación hidráulica, 4G CCO.", "Elevación y orientación en berma autorizada."),
        ("FRM_GENE", "Grupos Electrógenos", "Nivel diésel (>=75%), arranque (26V), 60Hz, 220V, ATS <10s.", "Conmutación ATS inmediata ante corte de red."),
        ("FRM_UPS", "Sistemas UPS", "Tensión banco 240V, 2.27V/celda, autonomía >=120 min.", "Transición limpia sin caída a modo batería."),
        ("FRM_SUBE", "Subestaciones Eléctricas", "Resistencia a tierra <=5.0 Ohm, aislamiento >=1000 MΩ.", "Tierra certificada con telurómetro calibrado."),
        ("FRM_FO", "Fibra Óptica (137 km)", "Atenuación <=0.25 dB/km, empalme <=0.05 dB, reflectancia.", "Margen óptico conforme para anillo ERPS."),
        ("FRM_SWIT", "Switches de Campo", "Enlace SFP (-15 a -8 dBm), 0 errores CRC, temp -20 a +70°C.", "Convergencia de anillo ITU-T <50 ms."),
        ("FRM_SWL3", "Switches Core Capa 3", "Uso CPU <35%, adyacencias BGP/OSPF, doble fuente AC.", "Cero interrupción de enrutamiento CCO."),
        ("FRM_ROUT", "Enrutadores de Borde", "Túneles VPN IPsec AES-256, throughput >=1 Gbps, firewall.", "Telemetría segura hacia Centro de Control."),
        ("FRM_SERV", "Servidores CCO", "Arreglo RAID 10 OK, disco <70%, uptime >=99.9%.", "Integridad de base de datos PostgreSQL."),
        ("FRM_NAS", "Almacenamiento NAS", "Retención video CCTV >=30 días, >=20 TB libres, SMART OK.", "Grabación continua de 30 días pliego ANI."),
        ("FRM_VW", "Video Wall CCO", "Matriz 4x2 activa, sin píxeles muertos, conmutación <1s.", "Despliegue simultáneo de cámaras y alarmas."),
        ("FRM_ETD", "Estación Tráfico ETD", "Conteo volumétrico, clasificación por ejes, sensor piezo.", "Precisión >=98% cotejada con aforo manual."),
        ("FRM_BASC", "Báscula Pesaje Estático", "Error pesaje bruto <=0.5%, simetría 4 celdas, plataforma.", "Certificado de pesas patrón vigente (SIC)."),
        ("FRM_BASD", "Báscula Dinámica WIM", "Sensor cuarzo, error peso eje <5% a 60 km/h, enlace CCO.", "Alerta en tiempo real de sobrepeso a peaje."),
        ("FRM_OCR", "Cámaras OCR Placas", "Tasa lectura >=96% placas, sincronía pulso infrarrojo.", "Identificación y vinculación al pesaje."),
        ("FRM_PJC", "Peaje: Carril", "Barrera <1.2s, semáforos LED carril, detector altura.", "Cero atascamiento en cabinas de peaje."),
        ("FRM_PJE", "Peaje: Telepeaje TAG", "Antena RFID 915 MHz, protocolo Colpass, lectura >=99.5%.", "Cobro electrónico sin detención (Colpass)."),
        ("FRM_PSEG", "Paso Seguro Peatonal", "Pulsador peatonal, balizas destello, sonorizador 85 dB.", "Activación inmediata de advertencia vial."),
        ("FRM_SGE", "Gestión Energía Peaje", "Factor de potencia >=0.95, banco condensadores, DPS.", "Cero penalización por reactiva ante operador."),
        ("FRM_SGM", "Gestión Ambiental", "Sensor visibilidad (niebla), pluviómetro, viento.", "Alertas meteorológicas automáticas en PMV."),
        ("FRM_SSA", "Señalización Acústica", "Sirenas túnel 110 dB a 3m, inteligibilidad STI >0.6.", "Evacuación audible en túneles del Sisga."),
        ("FRM_FIRE", "Extinción y Fuego", "Sensor térmico 68°C/88°C, pulsadores, mangueras >=6 bar.", "Activación de extractores en túnel."),
        ("FRM_PORT", "Portátiles Técnicos", "Batería >=4 horas, herramientas ITS, certificados VPN.", "Terminal de diagnóstico pericial en campo."),
        ("FRM_IMPR", "Impresoras Térmicas", "Corte limpio <0.5s, velocidad >=200 mm/s, sensor papel.", "Emisión continua de tiquetes de peaje.")
    ]

    for row_idx, r_data in enumerate(fichas_resumen):
        row = fichas_table.add_row()
        bg = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        for i, text in enumerate(r_data):
            cell = row.cells[i]
            cell.width = cols_f_w[i]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 50, 50, 50, 50)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            r.font.size = Pt(7.5)
            if i == 0:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # 4. VEREDICTO DE ACEPTACIÓN Y HOMOLOGACIÓN
    # -------------------------------------------------------------
    doc.add_heading("4. Veredicto de Homologación y Firmas de Aceptación", level=1)
    doc.add_paragraph(
        "Habiendo ejecutado y verificado satisfactoriamente los casos de prueba funcionales (UAT) y la carga técnica de las 27 fichas de inspección, las partes declaran que el Sistema de Gestión de Mantenimiento en Campo (SGMC v2) cumple a cabalidad con las especificaciones del Contrato de Concesión y el Apéndice Técnico 1 de la ANI, quedando HOMOLOGADO Y APTO PARA OPERACIÓN CONTINUA."
    ).paragraph_format.space_after = Pt(12)

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False

    sig_w = [Inches(3.2), Inches(3.2)]
    sig_content = [
        ("POR LA CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.:\n\n\n\n_____________________________________________\nIng. Diego Zúñiga\nCoordinador ITS / Director Técnico SGMC\nConcesión Transversal del Sisga S.A.S."),
        ("POR EL CONSORCIO DE INTERVENTORÍA:\n\n\n\n_____________________________________________\nIngeniero Especialista ITS / Auditor Contractual\nConsorcio Interventoría Sisga\nAgencia Nacional de Infraestructura (ANI)")
    ]

    for i, text in enumerate(sig_content):
        cell = sig_table.rows[0].cells[i]
        cell.width = sig_w[i]
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(text)
        r.font.size = Pt(8.5)

    doc.save(output_path)
    print(f"[OK] Documento de Pruebas UAT generado exitosamente en: {output_path}")

if __name__ == "__main__":
    out1 = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Manuales", "PLAN_DE_PRUEBAS_FUNCIONALES_Y_UAT_SGMC_V2.docx")
    generar_documento_word(out1)
