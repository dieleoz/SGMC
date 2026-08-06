# -*- coding: utf-8 -*-
"""Genera el documento Word de mesa de trabajo funcional del SGMC.

Salida: entregables/Definicion_Funcional_SGMC_Mesa_de_Trabajo.docx
Fuente de contenido: docs/DEFINICION_FUNCIONAL_MESA_DE_TRABAJO.md
Figuras: docs/images/fig_0*.png (generadas por generate_figuras.py)
Utilidades de composicion: scripts/_helpers_docx.py

Formato de respuesta: cada decision llega con una propuesta ya marcada [X].
El lider funcional confirma no escribiendo nada, o corrige marcando otra
opcion. Solo se pide redaccion libre donde no hay opcion posible (nombres,
fechas, listas propias). El documento viaja solo por correo, asi que las
instrucciones deben sostenerse sin presentacion en vivo.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import _helpers_docx as H
from _helpers_docx import (TINTA, GRIS, AZUL, ROJO, AMBAR, VERDE,
                           sombrear, bordes, barra_izquierda, p, titulo,
                           vinneta, bloque, tabla, figura, pie_de_pagina)

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
H.IMG = os.path.join(RAIZ, "docs", "images")
SALIDA = os.path.join(RAIZ, "entregables", "Definicion_Funcional_SGMC_Mesa_de_Trabajo.docx")

MARCA = "[X]"
VACIO = "[  ]"


# ------------------------------------------------------- bloques de respuesta
def opciones(doc, pregunta, items):
    """Bloque de opciones marcables. items = [(marcado_bool, texto), ...]"""
    if pregunta:
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(6)
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.left_indent = Cm(0.4)
        par.paragraph_format.keep_with_next = True
        r = par.add_run(pregunta)
        r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = TINTA
        r.font.name = "Segoe UI"
    for marcado, texto in items:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.9)
        par.paragraph_format.space_after = Pt(2)
        r = par.add_run((MARCA if marcado else VACIO) + "  ")
        r.font.size = Pt(10.5); r.font.bold = True
        r.font.color.rgb = VERDE if marcado else GRIS
        r.font.name = "Consolas"
        r = par.add_run(texto)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TINTA if marcado else GRIS
        r.font.bold = bool(marcado)
        r.font.name = "Segoe UI"


def campo(doc, etiqueta, largo=60):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.9)
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(etiqueta + "  ")
    r.font.size = Pt(10.5); r.font.color.rgb = TINTA; r.font.name = "Segoe UI"
    r = par.add_run("_" * largo)
    r.font.size = Pt(10.5); r.font.color.rgb = AMBAR; r.font.name = "Segoe UI"


def observacion(doc, lineas=2):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, "FFFDF5")
    bordes(c, color="B0700C", sz=6)
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run("SI DESEA MATIZAR O AGREGAR ALGO (opcional)")
    r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = AMBAR
    r.font.name = "Segoe UI"
    for _ in range(lineas):
        pp = c.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.space_before = Pt(0)
        rr = pp.add_run(" ")
        rr.font.size = Pt(10.5); rr.font.name = "Segoe UI"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# --------------------------------------------------------------- contenido
# Cada decision: bloques de opciones con la propuesta ya marcada, y campos
# de texto libre solo donde no cabe una opcion (nombres, fechas, listas).
DECISIONES = [
 dict(cod="D-01", tit="Coordenadas reales de los activos", bloque="A. Datos maestros", critica=True,
  enc="Los 34 activos del inventario tienen exactamente la misma coordenada: 4.728512, -74.114531. "
      "Ese punto está en Bogotá, no en el corredor.",
  imp="Toda la promesa de evidencia verificable del sistema descansa en el control GPS, y ese control compara la "
      "posición del técnico contra esa coordenada. Con el dato actual, un técnico parado frente a un poste SOS en "
      "Machetá no podrá cerrar nunca su mantenimiento, y en cambio cualquiera ubicado en ese punto de Bogotá "
      "validaría los 34 activos. La regla está bien programada; el dato la vuelve inútil.",
  grupos=[("¿De dónde tomamos las coordenadas reales?", [
             (1, "Recorrido de campo, capturando con el mismo celular que usará el técnico"),
             (0, "De un levantamiento topográfico o inventario georreferenciado que ya existe"),
             (0, "Otra fuente")]),
          ("¿Los 34 activos cargados son todo el inventario?", [
             (1, "Sí, son todos los activos que gestionará el sistema"),
             (0, "No, faltan activos por cargar")])],
  campos=["Responsable del levantamiento:", "Fecha comprometida:",
          "Si existe inventario georreferenciado, ¿dónde está?:",
          "Si faltan activos, ¿cuántos y de qué tipo?:"],
  desb="El control GPS, el piloto de campo y los criterios de aceptación CA-01.3 y CA-01.4."),

 dict(cod="D-02", tit="Radio de tolerancia del control GPS", bloque="A. Datos maestros",
  enc="El radio definido es 1,0 km, heredado de la especificación original.",
  imp="Un kilómetro es mucho para un poste SOS en vía abierta, donde 100 metros bastarían para probar presencia. Y "
      "puede ser poco para un activo lineal como un tramo de fibra óptica, donde el técnico puede intervenir "
      "legítimamente a varios kilómetros del punto registrado. Un radio único para 18 tipos de activo muy distintos "
      "es una simplificación que conviene revisar.",
  grupos=[("¿Qué radio aplicamos?", [
             (1, "Diferenciado: 200 m en activos puntuales (SOS, CCTV, paneles, básculas) y tratamiento aparte para fibra óptica"),
             (0, "Único de 1,0 km para todos los tipos, como está hoy"),
             (0, "Otro esquema")]),
          ("¿El kilómetro viene de alguna obligación?", [
             (0, "Sí, está en el contrato o lo exige interventoría"),
             (1, "No, fue un valor elegido por el equipo y se puede cambiar")])],
  campos=["Si es otro esquema, ¿cuál?:", "Si es contractual, ¿qué cláusula o documento?:"],
  desb="La configuración definitiva de la regla de validación."),

 dict(cod="D-03", tit="Qué significa la sede de un activo", bloque="A. Datos maestros", critica=True,
  enc="Los 11 usuarios están registrados en la sede 1, que es el CCO. Los 34 activos están repartidos entre las "
      "sedes 7 a 10, que son las unidades funcionales UF1 a UF4. La intersección es vacía.",
  imp="El filtro de seguridad descarga al celular del técnico solo los activos de su sede. Con los datos actuales, "
      "cada técnico descargaría cero activos y la aplicación quedaría vacía en sus manos el primer día de piloto. El "
      "problema de fondo es que la misma columna se usa con dos significados: para los usuarios es el sitio físico "
      "donde trabajan (CCO, peaje, báscula) y para los activos es el tramo del corredor donde están (UF1 a UF4).",
  grupos=[("¿Qué activos debe descargar un técnico en su celular?", [
             (1, "Los de las unidades funcionales que tenga asignadas, y puede tener varias"),
             (0, "Los de una sola unidad funcional"),
             (0, "Todos los del corredor, sin filtro")]),
          ("¿Qué debe ver un supervisor?", [
             (1, "Todo el corredor"),
             (0, "Solo su zona")]),
          ("¿Un técnico del CCO puede intervenir activos de cualquier unidad funcional?", [
             (1, "Sí"), (0, "No")])],
  campos=["Si aplica, ¿qué unidades funcionales atiende cada técnico?:"],
  desb="El filtro de seguridad y, con él, la descarga de datos al celular. Sin esto el piloto no arranca."),

 dict(cod="D-04", tit="Qué hacer cuando el GPS falla", bloque="B. Operación en campo",
  enc="El sistema captura la precisión del GPS pero no define qué hacer cuando esa precisión es mala. El plan "
      "original ya señalaba el riesgo de falsos negativos y proponía un mecanismo de excepción supervisada, que "
      "nunca se especificó.",
  imp="En túnel, en cañón o bajo copa densa, el celular puede reportar una posición con error de cientos de metros o "
      "no fijar nada. Si el sistema bloquea sin salida, el técnico pierde el trabajo de una hora y el sistema pierde "
      "su confianza el primer día. Si el sistema deja pasar todo, el control no sirve para nada. Este es el punto "
      "donde el control se gana o se pierde.",
  grupos=[("¿Qué hace el sistema cuando la precisión del GPS es insuficiente?", [
             (1, "Permite cerrar con excepción: exige motivo escrito y fotografía, marca el registro y avisa al supervisor"),
             (0, "Bloquea siempre: el técnico debe regresar cuando haya señal"),
             (0, "Permite cerrar solo con autorización previa por llamada al CCO")]),
          ("¿A partir de qué error del GPS se activa la excepción?", [
             (1, "Cuando la precisión reportada es peor que 50 metros"),
             (0, "Otro umbral")]),
          ("¿Las excepciones se reportan al supervisor?", [
             (1, "Sí, en un reporte semanal de excepciones"),
             (0, "No, basta con que queden registradas")])],
  campos=["Si es otro umbral, ¿cuál?:", "¿Qué porcentaje de excepciones consideraría un problema?:"],
  desb="La excepción E-02 del flujo de campo y buena parte de la aceptación del sistema por los técnicos."),

 dict(cod="D-05", tit="Interrupción a mitad de formulario", bloque="B. Operación en campo",
  enc="No hay definición de qué ocurre si el técnico cierra la aplicación, se le agota la batería o lo interrumpen a "
      "mitad del checklist.",
  imp="Una inspección de poste SOS son 15 preguntas más fotografías. Perder eso por una llamada entrante es la clase "
      "de fricción que hace que un técnico vuelva al papel.",
  grupos=[("¿Se conserva el trabajo a medio hacer?", [
             (1, "Sí, queda como borrador local y la inspección aparece en curso hasta cerrarse"),
             (0, "No, se descarta y el técnico reinicia")]),
          ("¿Hasta cuándo puede quedar abierta una inspección?", [
             (1, "Debe cerrarse dentro de la misma jornada"),
             (0, "Puede quedar abierta de un día para otro"),
             (0, "Otro plazo")]),
          ("¿Una inspección iniciada y no cerrada cuenta como incumplimiento?", [
             (0, "Sí, cuenta como incumplida"),
             (1, "No, pero se reporta aparte para seguimiento")])],
  campos=["Si es otro plazo, ¿cuál?:"],
  desb="La excepción E-03 y la definición del indicador de cumplimiento."),

 dict(cod="D-06", tit="Ciclo de vida de la orden de trabajo", bloque="B. Operación en campo",
  enc="Las órdenes registradas tienen estados Asignada, Cerrada y Suspendida, pero no hay regla que diga quién "
      "cambia cada estado ni bajo qué condición. Además, el único checklist existente apunta a una orden que no "
      "existe, señal de que hoy se puede registrar trabajo sin orden válida.",
  imp="Sin ciclo de vida definido no hay indicador de cumplimiento confiable: no se puede medir ejecutado contra "
      "programado si no está claro cuándo una orden cuenta como vencida, quién la suspende y si un trabajo puede "
      "existir sin orden.",
  grupos=[("¿Adoptamos los siete estados de la figura 4?", [
             (1, "Sí: Programada, Asignada, En ejecución, En revisión, Cerrada, Suspendida y Vencida"),
             (0, "No, preferimos otro conjunto de estados")]),
          ("¿Un técnico puede ejecutar un mantenimiento sin orden previa?", [
             (1, "Sí, lo levanta como novedad y el supervisor la convierte en orden"),
             (0, "No, siempre debe existir una orden asignada antes")]),
          ("¿Cuándo se considera vencida una orden?", [
             (1, "Al día siguiente de la fecha programada"),
             (0, "Al cierre de la semana"),
             (0, "Al cierre del mes")]),
          ("¿Se permite reasignar una orden a otro técnico?", [
             (1, "Sí, y queda traza de quién la reasignó"),
             (0, "No")])],
  campos=["Si prefiere otros estados, ¿cuáles?:", "¿Quién puede suspender una orden?:",
          "Motivos válidos de suspensión:"],
  desb="El indicador de cumplimiento, el flujo de asignación y la excepción E-06."),

 dict(cod="D-07", tit="Trabajo incompleto, segunda visita y devoluciones", bloque="B. Operación en campo",
  enc="El modelo prevé marcar que se requiere una segunda visita con su motivo, y la aprobación del supervisor, pero "
      "ninguno de los dos flujos está definido.",
  imp="El caso real más común no es el mantenimiento perfecto: es el técnico que llega y no puede terminar por falta "
      "de repuesto, por lluvia o por acceso cerrado. Si el sistema no modela ese caso, el técnico va a forzar un "
      "cierre falso o no va a registrar nada.",
  grupos=[("¿Qué pasa cuando el técnico no puede terminar?", [
             (1, "Cierre parcial con motivo tipificado, que genera una orden de seguimiento asociada a la original"),
             (0, "Se deja la orden abierta sin registrar nada"),
             (0, "Se suspende la orden y el supervisor decide")]),
          ("Motivos de trabajo incompleto que debe ofrecer el desplegable (marque los que apliquen)", [
             (1, "Falta de repuesto o material"),
             (1, "Condiciones climáticas"),
             (1, "Acceso restringido o activo inaccesible"),
             (1, "Riesgo para la seguridad del técnico"),
             (1, "Requiere personal o equipo especializado")]),
          ("Cuando el supervisor devuelve un mantenimiento, ¿qué ocurre?", [
             (1, "Se reabre el mismo registro al técnico, conservando la traza del rechazo"),
             (0, "El técnico crea un registro nuevo")]),
          ("¿La devolución notifica por correo al técnico?", [
             (1, "Sí"), (0, "No")])],
  campos=["Otros motivos a incluir:"],
  desb="La excepción E-04, el flujo de aprobación y el cálculo real de cumplimiento."),

 dict(cod="D-08", tit="Activos no inventariados y correctivo desde campo", bloque="B. Operación en campo",
  enc="No hay ruta para un activo que el técnico encuentra en vía y que no está en el sistema, ni para una falla "
      "detectada fuera de programación.",
  imp="Sin esa ruta, los hallazgos de campo se pierden o se gestionan por WhatsApp, que es exactamente lo que el "
      "sistema viene a reemplazar.",
  grupos=[("¿Puede el técnico reportar un activo que no está en el inventario?", [
             (1, "Sí, levanta una novedad con foto y coordenada; el supervisor decide si lo da de alta"),
             (0, "No, el inventario solo lo modifica el administrador")]),
          ("Cuando un activo queda fuera de servicio, ¿qué debe pasar?", [
             (1, "El sistema genera automáticamente una orden correctiva"),
             (0, "El supervisor la crea manualmente si lo considera"),
             (0, "El correctivo se gestiona por fuera del sistema")])],
  campos=[],
  desb="La excepción E-05 y el alcance de la alerta automática."),

 dict(cod="D-09", tit="Cuántos tipos de activo entran al primer sprint", bloque="C. Alcance de los formularios",
  critica=True,
  enc="Hay 18 tipos de activo y 18 formularios declarados, pero solo el de postes SOS tiene sus preguntas "
      "construidas: 15 preguntas. Los otros 17 están vacíos. Además, la columna que conecta cada tipo de activo con "
      "su formulario está vacía en los 18 tipos, de modo que hoy la aplicación no sabría qué checklist abrir ni "
      "siquiera para el SOS.",
  imp="Construir los 17 bancos restantes son del orden de 250 preguntas que alguien con criterio técnico debe "
      "redactar, con sus rangos, unidades y obligatoriedad. Es la tarea más grande que queda y es trabajo del equipo "
      "de la Concesión, no de configuración. Intentar construir los 18 a la vez es lo que hará que el proyecto se "
      "estanque otro mes.",
  grupos=[("¿Con qué tipos de activo arrancamos?", [
             (1, "Tres tipos: postes SOS, CCTV y paneles de mensaje variable"),
             (0, "Otros tres tipos"),
             (0, "Los 18 tipos a la vez")]),
          ("¿De dónde salen las preguntas de cada checklist?", [
             (1, "Se transcriben los formatos de inspección en papel que ya se usan"),
             (0, "Se redactan desde cero con el equipo técnico")]),
          ("¿Hay una exigencia contractual sobre qué debe contener una inspección?", [
             (0, "Sí, existe un requisito de interventoría o del contrato"),
             (1, "No, el contenido lo define el criterio técnico de la Concesión")])],
  campos=["Si son otros tres tipos, ¿cuáles?:", "Quién redacta y valida las preguntas:",
          "Disponibilidad semanal de esa persona:", "Si es contractual, ¿qué documento lo exige?:"],
  desb="El alcance real del primer sprint y su fecha. Es la decisión que más mueve el cronograma."),

 dict(cod="D-10", tit="Evidencia fotográfica y firmas", bloque="C. Alcance de los formularios",
  enc="El requerimiento habla de hasta 6 fotografías. El modelo de datos las soporta de dos maneras a la vez, "
      "incompatibles entre sí: dos campos fijos dentro del registro de mantenimiento y una tabla separada sin "
      "límite. Lo mismo ocurre con las firmas. Las tablas separadas están vacías.",
  imp="Si no se decide, el técnico terminará firmando dos veces y adjuntando fotos en dos lugares distintos, o peor, "
      "la evidencia quedará repartida y los reportes no la encontrarán completa.",
  grupos=[("¿Cuántas fotografías exige una inspección?", [
             (1, "Mínimo 3 obligatorias y hasta 6 en total"),
             (0, "Solo 2: una de inicio y una final"),
             (0, "Otra cantidad")]),
          ("¿Las fotografías son tipificadas o libres?", [
             (1, "Tipificadas: antes, después y novedad"),
             (0, "Libres, el técnico decide qué fotografía")]),
          ("¿Quién firma y dónde?", [
             (1, "El técnico firma en campo; el supervisor valida aprobando en el portal, sin firmar"),
             (0, "Firman ambos en campo, uno junto al otro")]),
          ("¿Qué valor tiene la firma?", [
             (0, "Contractual: es soporte frente a interventoría"),
             (1, "Control interno de la Concesión")])],
  campos=["Si es otra cantidad, ¿cuál?:"],
  desb="El diseño definitivo del formulario y el cumplimiento verificable de los requerimientos de evidencia."),

 dict(cod="D-11", tit="Trazabilidad histórica de las respuestas", bloque="C. Alcance de los formularios",
  enc="El detalle de las inspecciones guarda el texto de la pregunta, no su identificador. Si alguien reformula una "
      "pregunta, los registros anteriores dejan de ser comparables con los nuevos.",
  imp="Si el sistema debe mostrar la evolución de un activo en el tiempo, o demostrar ante interventoría que se "
      "aplicó el mismo criterio durante un periodo, esta trazabilidad es indispensable. Si el sistema solo debe dejar "
      "constancia de cada visita por separado, no lo es.",
  grupos=[("¿Necesita comparar la misma pregunta a lo largo del tiempo para un activo?", [
             (1, "Sí, es necesario ver la evolución del activo"),
             (0, "No, basta la constancia de cada visita por separado")]),
          ("¿El sistema debe poder reconstruir cómo era un formulario en una fecha pasada?", [
             (1, "Sí, se versionan los formularios al cambiarlos"),
             (0, "No hace falta")])],
  campos=[],
  desb="La comparabilidad histórica y los reportes de evolución."),

 dict(cod="D-12", tit="Qué reportes debe entregar el sistema", bloque="D. Reportes", critica=True,
  enc="Ningún reporte está definido. Se mencionan indicadores en documentos previos pero ninguno tiene fórmula, "
      "periodicidad ni destinatario.",
  imp="Este es el vacío más costoso del proyecto. Todo lo que el técnico captura en campo existe para producir algo, "
      "y ese algo nunca se especificó. Definir el reporte al final obliga casi siempre a volver atrás y capturar "
      "datos que no se pidieron. Definirlo ahora es lo que garantiza que el formulario pida lo correcto.",
  grupos=[("Marque los reportes que el sistema debe entregar. Los marcados son nuestra propuesta.", [
             (1, "Cumplimiento del plan de mantenimiento — mensual, a Supervisión y Dirección"),
             (1, "Activos fuera de servicio — diario, al CCO"),
             (1, "Hoja de vida del activo — a demanda, a Supervisión e Interventoría"),
             (1, "Certificado de mantenimiento por intervención — a Interventoría"),
             (0, "Productividad del técnico — semanal, a Supervisión"),
             (1, "Excepciones de GPS — semanal, a Supervisión")]),
          ("¿En qué formato deben entregarse?", [
             (1, "En pantalla, con exportación a Excel y PDF cuando se necesite enviar"),
             (0, "Solo en pantalla"),
             (0, "PDF automático por correo")]),
          ("¿Alguno se entrega a un tercero con formato obligatorio?", [
             (0, "Sí, hay un formato exigido por interventoría o la ANI"),
             (1, "No, el formato lo definimos nosotros")])],
  campos=["Otros reportes que necesite:", "Si hay formato obligatorio, ¿cuál y quién lo exige?:"],
  desb="El tablero y los entregables, y hacia atrás, la validación de que los formularios capturan lo necesario."),

 dict(cod="D-13", tit="Definición de los indicadores", bloque="D. Reportes",
  enc="Se habla de cumplimiento, disponibilidad y tiempo de atención sin fórmula acordada.",
  imp="La disponibilidad de activos puede significar tres cosas distintas según se mida por tiempo, por cantidad o "
      "ponderada por criticidad, y cada una da un número diferente ante la misma realidad. Si ese número va a un "
      "informe de interventoría, la definición debe estar acordada antes y no después.",
  grupos=[("¿Cómo se calcula el cumplimiento?", [
             (1, "Órdenes cerradas dentro de la fecha programada, sobre órdenes programadas"),
             (0, "Órdenes ejecutadas sobre programadas, sin importar la fecha")]),
          ("Una orden cerrada con excepción de GPS, ¿cuenta como cumplida?", [
             (1, "Sí, pero se reporta aparte en el informe de excepciones"),
             (0, "No cuenta como cumplida")]),
          ("¿Cómo se mide la disponibilidad de activos?", [
             (1, "Por tiempo: horas fuera de servicio sobre horas totales del periodo"),
             (0, "Por cantidad: activos operativos sobre activos totales"),
             (0, "Ponderada por criticidad del activo")]),
          ("¿Existe una meta contractual frente a la ANI que el sistema deba reportar?", [
             (0, "Sí, hay meta de disponibilidad o de tiempo de atención"),
             (0, "No"),
             (1, "Por confirmar con el área contractual")])],
  campos=["Si existe meta contractual, ¿cuál es y de qué documento sale?:"],
  desb="El tablero de indicadores y cualquier reporte a terceros."),

 dict(cod="D-14", tit="Usuarios, licenciamiento y gobierno del cambio", bloque="E. Gobierno",
  enc="Hay 11 usuarios registrados, dos de ellos inactivos. El presupuesto declarado en el plan original fue de 100 "
      "USD mensuales, y la plataforma se cobra por usuario activo. No hay definición de quién puede modificar la "
      "aplicación en producción.",
  imp="El costo escala con el número de usuarios, y sin un responsable de cambios cualquiera puede alterar un "
      "formulario en producción y romper la comparabilidad de los datos.",
  grupos=[("¿Interventoría o la ANI tendrán acceso al sistema?", [
             (0, "Sí, con perfil de consulta"),
             (1, "No, solo reciben reportes exportados")]),
          ("¿Quién puede modificar formularios y reglas en producción?", [
             (1, "Solo el administrador, con autorización escrita del responsable funcional"),
             (0, "El administrador, sin autorización previa")]),
          ("¿Cuánto tiempo se conserva la evidencia fotográfica?", [
             (1, "Cinco años"),
             (0, "Lo que dure la concesión"),
             (0, "Otro plazo")])],
  campos=["Número de técnicos que usarán el sistema en régimen:",
          "Responsable funcional que autoriza cambios:", "Si es otro plazo de retención, ¿cuál?:"],
  desb="El dimensionamiento de licencias y el procedimiento de control de cambios."),
]


def construir():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

    est = doc.styles["Normal"]
    est.font.name = "Segoe UI"
    est.font.size = Pt(10.5)
    est.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")

    pie_de_pagina(doc)

    # ---------------------------------------------------------- PORTADA
    p(doc, "CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.", size=10, bold=True, color=AZUL, space_after=2)
    p(doc, "Sistema de Gestión de Mantenimiento en Campo", size=10, color=GRIS, space_after=26)
    p(doc, "Definición funcional", size=30, bold=True, color=TINTA, space_after=0)
    p(doc, "y mesa de trabajo", size=30, bold=True, color=AZUL, space_after=16)
    p(doc, "Catorce decisiones para confirmar o corregir", size=13, color=GRIS, space_after=24)

    tabla(doc, ["", ""], [
        ["Dirigido a", "Líder funcional y equipo de Operaciones / ITS"],
        ["Elaborado por", "Equipo técnico del SGMC"],
        ["Fecha", "6 de agosto de 2026"],
        ["Versión", "2.0"],
        ["Estado del proyecto", "Sprint 0. Definición funcional en validación"],
        ["Documento previo", "Plan de Implementación del SGMC sobre Google AppSheet, julio de 2026"],
        ["Qué se espera de usted", "Confirmar o corregir 14 decisiones y devolver este archivo"],
        ["Tiempo estimado", "45 minutos"],
    ], anchos=[4.5, 12.0], size=10)

    p(doc, space_after=14)
    bloque(doc, "Cómo responder, en tres pasos.",
           "1) Lea la decisión. 2) Si está de acuerdo con la opción que aparece marcada, no escriba nada. "
           "3) Si no está de acuerdo, marque otra opción y, si hace falta, explique en el recuadro. "
           "Guarde el archivo y devuélvalo por correo. No necesita responder en documento aparte.",
           "1F5B99", AZUL, "E2ECF6")
    bloque(doc, "Lo que ya viene marcado es nuestra propuesta, no una decisión tomada.",
           "La marcamos para ahorrarle tiempo y para que quede claro qué haríamos si no recibimos respuesta. "
           "Cualquiera de ellas puede cambiarse. Lo que no se corrija se tomará como aceptado y quedará vinculante "
           "para la implementación.",
           "B0700C", AMBAR, "FDF3DE")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------- GUÍA
    titulo(doc, "Cómo está organizado este documento", 1)
    vinneta(doc, "Las secciones 1 a 4 explican el proyecto, qué encontramos al auditarlo y cómo entendemos que debe "
                 "funcionar el sistema. Son de lectura, para que las decisiones tengan contexto.")
    vinneta(doc, "La sección 5 son las catorce decisiones. Es donde usted responde.")
    vinneta(doc, "Las secciones 6 y 7 explican cómo sus respuestas se convierten en cronograma, y recogen el acta.")
    p(doc, space_after=4)
    p(doc, "Si dispone de poco tiempo, responda al menos las cuatro marcadas como ruta crítica: D-01, D-03, D-09 y "
           "D-12. Sin esas cuatro el proyecto no puede avanzar; las otras diez pueden quedar con la propuesta que "
           "traen marcada.", bold=True)

    # ---------------------------------------------------------- 1. CONTEXTO
    titulo(doc, "1. Por qué volvemos a la definición funcional", 1)
    p(doc, "El SGMC se planteó en julio de 2026 como un producto mínimo viable de 8 días sobre Google AppSheet, a "
           "partir de una especificación de requerimientos y un formulario de levantamiento diligenciado el 25 de "
           "julio. Se construyó, y hoy existe una aplicación publicada con un modelo de datos de 24 tablas, 34 "
           "activos catalogados y órdenes de trabajo registradas.")
    p(doc, "En el camino el alcance mutó. El modelo pasó de 17 a 24 tablas, aparecieron dos maneras distintas de "
           "definir los formularios de inspección, y las fotografías y firmas quedaron modeladas por duplicado. "
           "Ninguna de esas mutaciones pasó por una validación funcional.")
    p(doc, "Una auditoría del 6 de agosto de 2026 verificó ocho hallazgos bloqueantes leyendo directamente el archivo "
           "de producción. Al revisarlos uno a uno aparece un patrón que explica este documento:")
    bloque(doc, "",
           "La mayoría de los bloqueantes no son errores de programación. Son preguntas de negocio que nadie "
           "respondió. No se puede configurar el filtro de seguridad sin saber si la sede de un activo es la unidad "
           "funcional o el peaje. No se puede cerrar el requerimiento de evidencia fotográfica sin saber cuántas "
           "fotos exige realmente una inspección.",
           "1F5B99", AZUL, "E2ECF6")
    p(doc, "Por eso el paso siguiente no es seguir construyendo. Es definir, validar con usted, y recién entonces "
           "ejecutar.", space_before=6)

    # ---------------------------------------------------------- 2. HALLAZGOS
    titulo(doc, "2. Qué encontramos", 1)
    p(doc, "Verificado el 6 de agosto de 2026 leyendo el archivo maestro de producción. Lo que está construido y "
           "funcionando: el modelo de 24 tablas, los catálogos poblados con 34 activos y sus códigos QR, 18 tipos de "
           "activo, 10 sedes, 11 usuarios, 6 órdenes de trabajo, un checklist de inspección con su detalle, el banco "
           "de preguntas del formulario de postes SOS y la aplicación publicada con sus vistas móviles y web.")
    p(doc, "Lo que está abierto:")
    tabla(doc, ["#", "Hallazgo", "Decisión"], [
        ["B-01", "Los 34 activos comparten una sola coordenada, situada en Bogotá y no en el corredor. El control GPS "
                 "es inoperante hasta levantar las coordenadas reales", "D-01"],
        ["B-02", "La columna que conecta cada tipo de activo con su formulario está vacía en los 18 tipos", "D-09"],
        ["B-03", "Todos los usuarios están en la sede 1 y todos los activos en las sedes 7 a 10. El filtro de "
                 "seguridad dejaría a cada técnico sin activos", "D-03"],
        ["B-04", "Solo 1 de 18 formularios tiene su banco de preguntas construido", "D-09"],
        ["B-05", "El único checklist existente referencia una orden de trabajo que no existe", "D-06"],
        ["B-06", "Fotografías, firmas y posición GPS están modelados dos veces", "D-10"],
        ["B-07", "La tabla de mantenimientos está vacía: ningún mantenimiento se ha ejecutado de extremo a extremo",
         "Todas"],
        ["B-08", "El detalle de checklist guarda las preguntas como texto libre, sin trazabilidad", "D-11"],
    ], anchos=[1.6, 11.9, 3.0])
    p(doc, "El hallazgo B-01 merece verse en detalle, porque es el que invalida la promesa central del sistema.",
      space_before=6)
    figura(doc, "fig_03_coordenadas.png",
           "Figura 3. Los 34 activos tienen registrada la misma coordenada, ubicada en Bogotá y no en el corredor.")

    # ---------------------------------------------------------- 3. ACTORES
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "3. Actores y alcance", 1)
    p(doc, "Cuatro roles operan el sistema. Confirme que la descripción corresponde a la organización real y que no "
           "falta ningún actor.")
    figura(doc, "fig_01_actores.png", "Figura 1. Actores del sistema, ámbito de trabajo y restricción de datos.")

    titulo(doc, "Alcance propuesto para el primer sprint productivo", 2)
    p(doc, "Entra: mantenimiento preventivo programado, sobre los tipos de activo que se prioricen, ejecutado por "
           "técnico con evidencia fotográfica, firma y validación GPS, con asignación y aprobación del supervisor, y "
           "un tablero básico de cumplimiento.")
    p(doc, "No entra por ahora: generación automática de órdenes por frecuencia, integración con Power BI o mesas de "
           "ayuda, gestión de repuestos e inventario, y firma de interventoría. Las decisiones D-08, D-09 y D-12 "
           "pueden mover este corte.")

    # ---------------------------------------------------------- 4. FLUJOS
    titulo(doc, "4. Flujos funcionales", 1)
    p(doc, "Cada flujo se expresa con precondiciones, pasos, excepciones y criterios de aceptación verificables, de "
           "modo que sirva a la vez como definición funcional y como base de las pruebas de aceptación de usuario.")

    titulo(doc, "CU-01. El técnico ejecuta un mantenimiento en campo", 2, color=AZUL)
    p(doc, "Es el flujo que justifica el sistema. Los pasos en ámbar dependen de una decisión pendiente.")
    figura(doc, "fig_02_flujo_tecnico.png",
           "Figura 2. Ciclo del técnico en campo y excepciones sin respuesta definida.")
    p(doc, "Criterios de aceptación", bold=True, space_before=6, space_after=4)
    tabla(doc, ["Código", "El sistema se considera conforme cuando"], [
        ["CA-01.1", "El registro se completa de principio a fin sin señal celular en ningún momento"],
        ["CA-01.2", "Al recuperar señal, el registro aparece en el sistema sin acción del técnico"],
        ["CA-01.3", "Estando a más de la distancia definida del activo, el sistema impide guardar"],
        ["CA-01.4", "Estando junto al activo, el sistema permite guardar"],
        ["CA-01.5", "El checklist que abre corresponde al tipo del activo escaneado"],
        ["CA-01.6", "Las fotografías quedan asociadas al mantenimiento y comprimidas"],
        ["CA-01.7", "Un mantenimiento sin firma no puede cerrarse"],
    ], anchos=[2.6, 13.9])
    bloque(doc, "Advertencia sobre el estado actual.",
           "Este flujo nunca se ha ejecutado completo: la tabla de mantenimientos está vacía. Y como los 34 activos "
           "comparten una coordenada en Bogotá, hoy fallan a la vez CA-01.3 y CA-01.4.",
           "A82D2D", ROJO, "FBE9E9")

    titulo(doc, "CU-02. El supervisor programa y asigna una orden de trabajo", 2, color=AZUL)
    p(doc, "El supervisor crea la orden en el portal seleccionando activo, técnico, fecha programada y tipo de "
           "trabajo. El sistema asigna el número, un bot notifica por correo al técnico y la orden aparece en su "
           "celular al sincronizar. Los estados y quién los mueve se deciden en D-06.")
    figura(doc, "fig_04_ciclo_ot.png", "Figura 4. Ciclo de vida propuesto para la orden de trabajo.")

    titulo(doc, "CU-03. El supervisor revisa y aprueba", 2, color=AZUL)
    p(doc, "Al sincronizar el técnico, el mantenimiento aparece en el portal. El supervisor revisa checklist, "
           "fotografías, firma y coordenada de cierre, y aprueba o devuelve con observación. Al aprobar, la orden se "
           "cierra y el indicador de cumplimiento se actualiza. El flujo de devolución se decide en D-07.")

    titulo(doc, "CU-04. Alerta automática por activo fuera de servicio", 2, color=AZUL)
    p(doc, "Cuando el técnico cierra un mantenimiento con estado final Fuera de servicio, al sincronizar se dispara "
           "un bot que genera un informe en PDF con los datos de la falla, la ubicación y las fotografías, y lo envía "
           "por correo prioritario al CCO y al supervisor de la zona. Si eso genera además una orden correctiva se "
           "decide en D-08.")

    titulo(doc, "CU-05. El administrador mantiene el inventario y los formularios", 2, color=AZUL)
    p(doc, "Da de alta activos con su código, tipo, sede, punto de referencia, calzada, sentido, coordenada real, "
           "frecuencia y código QR. Define o ajusta el formulario de inspección de cada tipo de activo y publica el "
           "cambio, que los técnicos reciben en la siguiente sincronización. La trazabilidad al modificar preguntas "
           "se decide en D-11.")

    titulo(doc, "CU-06. Reportes y tablero", 2, color=ROJO)
    p(doc, "Este caso de uso no está definido, y es el vacío más grande del proyecto: el sistema captura datos desde "
           "julio pero nadie ha especificado qué debe producir con ellos. Se define en D-12 y D-13.")

    # ---------------------------------------------------------- 5. DECISIONES
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "5. Las catorce decisiones", 1)
    p(doc, "Recuerde: si está de acuerdo con lo marcado, no escriba nada. Marque solo lo que quiera cambiar.")
    p(doc, "Las cuatro decisiones señaladas como ruta crítica son las que fijan el cronograma. Si solo puede "
           "responder algunas, que sean esas.", color=GRIS, italic=True)

    bloque_actual = None
    for d in DECISIONES:
        if d["bloque"] != bloque_actual:
            bloque_actual = d["bloque"]
            titulo(doc, "Bloque " + bloque_actual, 2, color=GRIS, size=11)

        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(14)
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.keep_with_next = True
        r = par.add_run(d["cod"] + ".  ")
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = AZUL; r.font.name = "Segoe UI"
        r = par.add_run(d["tit"])
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = TINTA; r.font.name = "Segoe UI"
        if d.get("critica"):
            r = par.add_run("     RUTA CRÍTICA")
            r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = ROJO; r.font.name = "Segoe UI"

        p(doc, "Qué encontramos", size=9.5, bold=True, color=GRIS, space_before=6, space_after=2)
        p(doc, d["enc"], space_after=6)
        p(doc, "Por qué importa", size=9.5, bold=True, color=GRIS, space_after=2)
        p(doc, d["imp"], space_after=6)

        p(doc, "Nuestra propuesta. Confirme o corrija.", size=9.5, bold=True, color=AMBAR, space_after=2)
        for pregunta, items in d["grupos"]:
            opciones(doc, pregunta, items)
        if d["campos"]:
            p(doc, space_after=2)
            for c in d["campos"]:
                campo(doc, c)
        observacion(doc, lineas=2)
        p(doc, "Desbloquea: " + d["desb"], size=9, italic=True, color=GRIS, space_after=10)

    # ---------------------------------------------------------- 6. ROADMAP
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "6. De sus respuestas al cronograma", 1)
    p(doc, "El cronograma no está escrito porque depende de estas respuestas. Este es el mapa de dependencias.")
    figura(doc, "fig_05_ruta_critica.png", "Figura 5. Secuencia de ejecución y ruta crítica del proyecto.")
    p(doc, "La ruta crítica son D-01 y D-09. Levantar las coordenadas de los 34 activos y redactar los bancos de "
           "preguntas son trabajo del equipo de la Concesión, no de configuración, y ambas se miden en semanas. El "
           "cronograma completo lo fijan esas dos y no las demás.", bold=True, space_before=4)

    titulo(doc, "Orden de ejecución una vez recibidas sus respuestas", 2)
    for i, (t, s) in enumerate([
        ("Definición", "Estas respuestas, consolidadas en un acta de decisiones."),
        ("Datos", "Coordenadas reales, mapeo de formulario por tipo de activo, realineación de sedes y construcción "
                  "de los bancos de preguntas priorizados."),
        ("Configuración", "Reglas de validación, filtro de seguridad, formularios, bots de notificación y reportes."),
        ("Prueba controlada", "Un mantenimiento completo de extremo a extremo ejecutado por una persona, primero con "
                              "señal y luego en modo avión."),
        ("Piloto", "Los 10 celulares en vía, solo después de que la prueba controlada haya funcionado."),
        ("Producción y evolución", "Resto de tipos de activo, integraciones y respaldo automático."),
    ], 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(4)
        r = par.add_run(f"{i}. {t}. ")
        r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = AZUL; r.font.name = "Segoe UI"
        r = par.add_run(s)
        r.font.size = Pt(10.5); r.font.color.rgb = TINTA; r.font.name = "Segoe UI"

    # ---------------------------------------------------------- 7. CIERRE
    titulo(doc, "7. Cierre", 1)
    p(doc, "Para diligenciar al devolver el documento.")
    tabla(doc, ["Campo", "Registro"], [
        ["Diligenciado por", ""],
        ["Cargo", ""],
        ["Fecha", ""],
        ["Decisiones que dejó como estaban (aceptadas)", ""],
        ["Decisiones que corrigió", ""],
        ["Decisiones que no puede responder y por qué", ""],
        ["Responsable de D-01, coordenadas, y fecha comprometida", ""],
        ["Responsable de D-09, bancos de preguntas, y fecha comprometida", ""],
        ["¿Requiere reunión para alguna decisión? ¿Cuáles?", ""],
    ], anchos=[7.0, 9.5], size=10)

    p(doc, space_after=20)
    t = doc.add_table(rows=1, cols=2)
    for i, et in enumerate(["Líder funcional", "Por el equipo técnico"]):
        c = t.cell(0, i)
        c.paragraphs[0].text = ""
        pp = c.paragraphs[0]
        r = pp.add_run("\n\n____________________________________")
        r.font.size = Pt(10.5); r.font.name = "Segoe UI"; r.font.color.rgb = TINTA
        pp2 = c.add_paragraph()
        r = pp2.add_run(et + "\nNombre y fecha")
        r.font.size = Pt(9); r.font.color.rgb = GRIS; r.font.name = "Segoe UI"

    p(doc, space_after=14)
    p(doc, "Las decisiones registradas en este documento son la base del cronograma de implementación y del alcance "
           "del sprint. Lo que no se corrija se toma como aceptado.", size=9, italic=True, color=GRIS)

    doc.save(SALIDA)
    print("Documento generado:", SALIDA)


if __name__ == "__main__":
    construir()
