# -*- coding: utf-8 -*-
"""
Generador del Manual Funcional y Guía de Operación por Rol en formato Microsoft Word (.docx)
para la Concesión Transversal del Sisga S.A.S. (SGMC v2).
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Establece color de fondo a una celda de tabla."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece márgenes internos a una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generar_manual_word(output_path):
    doc = Document()

    # Configuración de márgenes (1 pulgada / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Estilos de Fuente Base
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(15, 23, 42) # Slate-900

    # -------------------------------------------------------------
    # PORTADA / ENCABEZADO PRINCIPAL
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    run_concesion = title_p.add_run("CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.")
    run_concesion.bold = True
    run_concesion.font.size = Pt(13)
    run_concesion.font.color.rgb = RGBColor(4, 120, 87) # Emerald-700

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("SISTEMA DE GESTIÓN DE MANTENIMIENTO EN CAMPO (SGMC v2)\nMANUAL FUNCIONAL Y GUÍA DE OPERACIÓN POR ROL")
    run_sub.bold = True
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(15, 23, 42)

    # Fila de Metadatos del Documento
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    headers_meta = [
        ("Versión:", "2.0 (PostGIS)"),
        ("Corredor:", "137 km (Sisga - Aguaclara)"),
        ("Censo Activos:", "368 Equipos"),
        ("Emisión:", "Agosto 2026")
    ]
    
    for i, (label, val) in enumerate(headers_meta):
        cell = meta_table.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"{label} ")
        r1.bold = True
        r1.font.size = Pt(8.5)
        r2 = p.add_run(val)
        r2.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 1: PROPÓSITO Y ALCANCE
    # -------------------------------------------------------------
    h1 = doc.add_heading("1. Propósito y Alcance del SGMC v2", level=1)
    h1.style.font.color.rgb = RGBColor(15, 23, 42)
    
    p1 = doc.add_paragraph(
        "El Sistema de Gestión de Mantenimiento en Campo (SGMC v2) es la plataforma oficial de la Concesión Transversal del Sisga para la administración, planificación, ejecución técnica y certificación pericial del mantenimiento preventivo y correctivo sobre los 368 activos de infraestructura distribuidos en las cuatro Unidades Funcionales del corredor vial:"
    )
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(6)

    ufs = [
        ("UF1 — Sisga / Guateque (PK 00+000 al PK 49+000):", " 146 activos (Postes SOS, Cámaras CCTV, Paneles PMV, Subestaciones)."),
        ("UF2 — Guateque / Macanal (PK 49+000 al PK 72+000):", " 53 activos (Túneles, Estaciones de Peaje, Cámaras de Monitoreo)."),
        ("UF3 — Macanal / Santa María (PK 72+000 al PK 105+000):", " 45 activos (Postes SOS, Radares de Velocidad, Luminarias)."),
        ("UF4 — Santa María / Aguaclara (PK 105+000 al PK 137+000):", " 124 activos (Sistemas de Comunicación, Estaciones Meteorológicas).")
    ]
    for uf_title, uf_desc in ufs:
        p_uf = doc.add_paragraph(style='List Bullet')
        p_uf.paragraph_format.space_after = Pt(3)
        r_t = p_uf.add_run(uf_title)
        r_t.bold = True
        p_uf.add_run(uf_desc)

    doc.add_paragraph(
        "El sistema garantiza trazabilidad jurídica e inmutabilidad pericial ante la Agencia Nacional de Infraestructura (ANI) y la Interventoría mediante validación satelital fail-closed, evidencias fotográficas WebP en almacenamiento S3, firma digital y cálculo automatizado de la Disponibilidad Contractual (Di >= 98.5%)."
    ).paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 2: MATRIZ DE ROLES Y RESPONSABILIDADES
    # -------------------------------------------------------------
    doc.add_heading("2. Matriz de Roles y Responsabilidades", level=1)
    
    roles_table = doc.add_table(rows=1, cols=4)
    roles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roles_table.autofit = False

    cols_w = [Inches(1.2), Inches(1.5), Inches(1.3), Inches(2.5)]
    headers_roles = ["Rol", "Perfiles", "Interfaz URL", "Responsabilidades Funcionales"]
    
    for i, h_text in enumerate(headers_roles):
        cell = roles_table.rows[0].cells[i]
        cell.width = cols_w[i]
        set_cell_background(cell, "0F172A") # Slate-900
        set_cell_margins(cell, 100, 100, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    roles_data = [
        ("ROL-03\nTécnico de Campo", "Iván Salcedo\nLuis Gacha", "/tecnico\n(PWA Móvil)", "• Ejecutar inspecciones en modo 100% offline.\n• Diligenciar checklists dinámicos por tipo de activo.\n• Captura de fotos WebP con georreferenciación y firmas.\n• Reportar novedades de vía en carretera (/novedades)."),
        ("ROL-02\nSupervisor de Zona", "Fernand Bolívar", "/supervisor\n/planes\n(Web / Tablet)", "• Auditar órdenes ejecutadas y tolerancias GPS.\n• Aprobar o rechazar mantenimientos con observaciones.\n• Generar Fichas Técnicas Periciales en PDF.\n• Programar planes preventivos mensuales por UF (/planes)."),
        ("ROL-01\nDirector Técnico / CCO", "Diego Zúñiga\nOperadores CCO", "/\n/activos\n(Web Desktop)", "• Monitoreo general del corredor vial (137 km).\n• Administrar el catálogo maestro de 368 activos.\n• Supervisar asignaciones de técnicos por zona."),
        ("ROL-04\nInterventoría / ANI", "Consorcio Interventoría Sisga", "/reportes\n(Web Desktop)", "• Auditar Disponibilidad Contractual (Di >= 98.5%).\n• Consultar Parte Diario de Operaciones del CCO.\n• Descargar el Informe Oficial en PDF para la ANI.")
    ]

    for row_idx, r_data in enumerate(roles_data):
        row = roles_table.add_row()
        bg = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        for i, text in enumerate(r_data):
            cell = row.cells[i]
            cell.width = cols_w[i]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 3: GUÍA OPERATIVA ROL TÉCNICO DE CAMPO
    # -------------------------------------------------------------
    doc.add_heading("3. Manual de Operación: Rol Técnico de Campo (/tecnico)", level=1)
    
    doc.add_heading("3.1. Instalación de la Aplicación en Celular o Tablet (PWA)", level=2)
    p_inst = doc.add_paragraph(
        "1. Abra el navegador Google Chrome (Android) o Safari (iOS) en el celular.\n"
        "2. Ingrese a la dirección: https://sisga-2.vercel.app/tecnico\n"
        "3. Presione el botón 'Instalar App' en el banner superior (o en el menú del navegador: 'Instalar aplicación' / 'Agregar a inicio').\n"
        "4. La aplicación quedará instalada en la pantalla del celular como una App nativa, con rendimiento optimizado y caché fuera de línea."
    )
    p_inst.paragraph_format.line_spacing = 1.15
    p_inst.paragraph_format.space_after = Pt(6)

    doc.add_heading("3.2. Operación Offline en Túneles y Zonas sin Señal", level=2)
    doc.add_paragraph(
        "• Almacenamiento Local (IndexedDB): Todas las órdenes asignadas y formularios se guardan en la memoria local del dispositivo.\n"
        "• Indicador de Conectividad: Si entra a un túnel o sector sin cobertura celular, la cabecera mostrará 'Modo Offline Activo (Túneles/Vía)'.\n"
        "• Contador de Pendientes: Si tiene inspecciones guardadas sin sincronizar, verá un badge ámbar (ej. '⚠️ 2 pendientes').\n"
        "• Sincronización Automática: Al salir a una zona con señal 4G o WiFi, el sistema sube automáticamente los registros a Supabase sin perder información."
    ).paragraph_format.space_after = Pt(6)

    doc.add_heading("3.3. Ciclo Paso a Paso de una Inspección en Sitio", level=2)
    doc.add_paragraph(
        "1. Seleccionar la Orden Asignada: Toque la tarjeta del activo a intervenir (ej. Poste SOS PK 14+200).\n"
        "2. Captura Satelital GPS (Geofencing):\n"
        "   - Presione 'Capturar Coordenadas GPS'. El sistema valida la presencia dentro del radio de tolerancia (ej. 50 m).\n"
        "   - Cierre con Excepción (Túneles): Si no hay señal de satélite por estar bajo techo o en túnel, marque 'Activar Cierre con Excepción Manual' e indique el motivo. El sistema registrará la labor de forma limpia sin inventar coordenadas.\n"
        "3. Diligenciamiento de Checklist Dinámico: Responda las preguntas técnicas por sección (Conforme/No conforme, listas de valores, voltajes numéricos en voltios, etc.).\n"
        "4. Fotografías WebP: Capture las fotos de evidencia. La cámara estampará automáticamente fecha, hora y ubicación en el pie de foto con compresión WebP (<150 KB).\n"
        "5. Firma Digital Manuscrita: Dibuje su firma con el dedo en el recuadro táctil.\n"
        "6. Guardar Mantenimiento: Presione 'Guardar y Cerrar Mantenimiento'. El botón se bloqueará contra doble pulsación y asentará la orden para supervisión."
    ).paragraph_format.space_after = Pt(6)

    doc.add_heading("3.4. Reporte Rápido de Novedades en Ruta (/novedades)", level=2)
    doc.add_paragraph(
        "Si durante el recorrido en carretera detecta un cable cortado, un poste chocado o un daño imprevisto:\n"
        "1. Presione el botón 'Reportar Novedad en Ruta' en la cabecera de la PWA.\n"
        "2. Seleccione el tipo de falla, tome la foto de evidencia y capture el GPS.\n"
        "3. Al enviar, el sistema genera automáticamente una Orden de Trabajo Correctiva (OT-CORR) asignada a la zona correspondiente."
    ).paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 4: GUÍA OPERATIVA ROL SUPERVISOR DE MANTENIMIENTO
    # -------------------------------------------------------------
    doc.add_heading("4. Manual de Operación: Rol Supervisor de Mantenimiento (/supervisor y /planes)", level=1)
    
    doc.add_heading("4.1. Bandeja de Auditoría y Certificación (/supervisor)", level=2)
    doc.add_paragraph(
        "1. Ingrese a https://sisga-2.vercel.app/supervisor desde su computador o tablet.\n"
        "2. Filtre por Unidad Funcional (UF1 a UF4), estado de la orden ('En revision' o 'Cerrada') o técnico asignado.\n"
        "3. Presione 'Auditar Evidencias' en la orden que desea revisar."
    ).paragraph_format.space_after = Pt(6)

    doc.add_heading("4.2. Modal Pericial de Supervisión y Aprobación", level=2)
    doc.add_paragraph(
        "• Auditoría Geográfica: Verifique en el mapa interactivo la distancia real a la que el técnico cerró la orden frente al PK oficial del activo.\n"
        "• Inspección Visual: Examine las fotografías WebP en alta resolución almacenadas en Supabase Storage y la firma manuscrita.\n"
        "• Aprobación / Rechazo:\n"
        "   - Aprobar y Certificar: La orden pasa a estado 'Cerrada' y se incorpora al cálculo de disponibilidad contractual.\n"
        "   - Rechazar / Solicitar Corrección: Permite devolver la orden al técnico con observaciones técnicas para su subsanación.\n"
        "• Descarga de Ficha PDF: Presione 'Descargar Ficha PDF (Interventoría)' para generar el reporte pericial individual con membrete oficial."
    ).paragraph_format.space_after = Pt(6)

    doc.add_heading("4.3. Generador Automático de Planes Preventivos Mensuales (/planes)", level=2)
    doc.add_paragraph(
        "1. Ingrese a https://sisga-2.vercel.app/planes.\n"
        "2. Seleccione el ciclo mensual (ej. Septiembre 2026) y la Unidad Funcional (o 'Todas las UFs').\n"
        "3. Presione 'Generar OTs del Mes'.\n"
        "4. El sistema ejecuta el procedimiento masivo en PostgreSQL, programando los 368 activos en PLA_PlanMantenimiento y generando el lote de órdenes en OT_OrdenesTrabajo para cada cuadrilla."
    ).paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 5: GUÍA OPERATIVA ROL DIRECTOR TÉCNICO Y CCO
    # -------------------------------------------------------------
    doc.add_heading("5. Manual de Operación: Rol Director Técnico y Centro de Control (/ y /activos)", level=1)
    doc.add_paragraph(
        "• Centro de Control Operativo (/): Tablero global con métricas en tiempo real de avance preventivo y correctivo a lo largo de los 137 km de corredor vial.\n"
        "• Inventario Maestro de Activos (/activos): Búsqueda, consulta y georreferenciación del censo completo de 368 activos por código (ACT-0001), PK, Unidad Funcional y tipo de subsistema."
    ).paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECCIÓN 6: GUÍA OPERATIVA ROL INTERVENTORÍA Y AUDITORÍA ANI
    # -------------------------------------------------------------
    doc.add_heading("6. Manual de Operación: Rol Interventoría y Auditoría ANI (/reportes)", level=1)
    
    doc.add_heading("6.1. Tablero de Disponibilidad Contractual (Di)", level=2)
    doc.add_paragraph(
        "El sistema calcula automáticamente la fórmula oficial del Apéndice Técnico 1 de la ANI para cada uno de los subsistemas y Unidades Funcionales:\n\n"
        "    Di = [ 1 - ( Horas Indisponibles Totales / ( N° Activos × 720 h ) ) ] × 100%\n\n"
        "• Semáforo de Conformidad ANI:\n"
        "   - Di >= 98.5% -> CONFORME (Retribución contractual al 100%).\n"
        "   - Di < 98.5% -> NO CONFORME (Genera deducción de disponibilidad contractual)."
    ).paragraph_format.space_after = Pt(6)

    doc.add_heading("6.2. Parte Diario de Operaciones del CCO", level=2)
    doc.add_paragraph(
        "Permite auditar el volumen diario de órdenes programadas, mantenimientos ejecutados, cierres con excepción satelital y novedades atendidas en carretera."
    ).paragraph_format.space_after = Pt(6)

    doc.add_heading("6.3. Emisión del Informe Oficial en PDF para Radicación ante la ANI", level=2)
    doc.add_paragraph(
        "1. En la vista /reportes, presione el botón verde 'Descargar Informe PDF (ANI)'.\n"
        "2. El sistema compila instantáneamente el balance mensual con membrete de la Concesión, código INF-DISP-YYYYMM, matriz de los 27 subsistemas y bloques de firma para el Director Técnico y el Ingeniero Residente de Interventoría.\n"
        "3. El documento se abre listo para impresión o guardado en PDF para archivo contractual."
    ).paragraph_format.space_after = Pt(16)

    # -------------------------------------------------------------
    # BLOQUE DE FIRMAS Y CERTIFICACIÓN
    # -------------------------------------------------------------
    doc.add_heading("7. Certificación de Aprobación del Manual", level=1)
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False

    sig_w = [Inches(3.2), Inches(3.2)]
    sig_content = [
        ("POR LA CONCESIÓN TRANSVERSAL DEL SISGA:\n\n\n\n_____________________________________________\nIng. Diego Zúñiga\nCoordinador ITS / Director Técnico SGMC\nConcesión Transversal del Sisga S.A.S."),
        ("POR EL CONSORCIO DE INTERVENTORÍA:\n\n\n\n_____________________________________________\nIngeniero Especialista ITS / Auditor Contractual\nConsorcio Interventoría Sisga\nAgencia Nacional de Infraestructura (ANI)")
    ]

    for i, text in enumerate(sig_content):
        cell = sig_table.rows[0].cells[i]
        cell.width = sig_w[i]
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, 120, 120, 120, 120)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(text)
        r.font.size = Pt(9)

    doc.save(output_path)
    print(f"[OK] Documento Word generado exitosamente en: {output_path}")

if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Manuales", "MANUAL_FUNCIONAL_SGMC_V2.docx")
    generar_manual_word(out)
