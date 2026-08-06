import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

print("Generando Manual_de_Usuario_SGMC_Con_Diagramas.docx...")

doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Cover Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("MANUAL DE USUARIO Y GUÍA OPERATIVA ILUSTRADA\nSISTEMA DE GESTIÓN DE MANTENIMIENTO EN CAMPO (SGMC)")
run_title.bold = True
run_title.font.size = Pt(18)
run_title.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Concesión Transversal del Sisga S.A.S. | Plataforma: Google AppSheet\nVersión 1.0 (Manual Inicial para Validación Funcional) | Fecha: Agosto de 2026")
run_sub.font.size = Pt(11)
run_sub.font.italic = True

doc.add_paragraph() # Spacing

# 1. PRESENTACIÓN Y ALCANCE
doc.add_heading("1. Presentación y Perfiles de Usuario", level=1)
doc.add_paragraph(
    "El presente manual constituye la guía oficial de usuario para la operación del Sistema de Gestión de "
    "Mantenimiento en Campo (SGMC) de la Concesión Transversal del Sisga S.A.S. La plataforma permite la digitalización "
    "de mantenimientos preventivos y correctivos sobre la infraestructura vial (Postes SOS, CCTV, PMVF, PMVM, Sensores Ambientales, Básculas, Generadores y TI)."
)

# Table for Profiles
t_prof = doc.add_table(rows=1, cols=3)
t_prof.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_p = t_prof.rows[0].cells
hdr_p[0].text = "Perfil / Rol"
hdr_p[1].text = "Entorno de Uso"
hdr_p[2].text = "Funciones Principales"

profiles_data = [
    ("Técnico de Campo", "App Móvil (Smartphone Android/iOS)", "Inspección offline, escáner QR, fotos, firmas y geocierre GPS"),
    ("Supervisor CCO", "Portal Web (Computador / Tablet)", "Programación de OTs, revisión de alertas y tablero KPI"),
    ("Consulta / Interventoría", "Portal Web (Computador)", "Auditoría de intervenciones y exportación de reportes PDF/Excel"),
    ("Administrador de Sistema", "Portal Web (Computador)", "Gestión de usuarios, sedes, activos y plantillas de checklist")
]

for p in profiles_data:
    c = t_prof.add_row().cells
    c[0].text = p[0]
    c[1].text = p[1]
    c[2].text = p[2]

doc.add_paragraph() # Spacing

# 2. GUÍA ILUSTRADA PARA TÉCNICOS DE CAMPO
doc.add_heading("2. Guía Ilustrada para Técnicos de Campo (App Móvil)", level=1)

doc.add_heading("2.1 Paso 1: Inicio de Sesión y Descarga Inicial", level=2)
doc.add_paragraph(
    "1. Abra la aplicación AppSheet en su dispositivo móvil.\n"
    "2. Seleccione 'Sign in with Microsoft' (o Google) e ingrese su correo corporativo M365.\n"
    "3. El servidor evaluará su SedeID asignada (ej. Peaje Machetá, Sutatenza, San Luis de Gaceno) y descargará a la caché local únicamente los activos y OTs de su zona."
)

doc.add_heading("2.2 Paso 2: Navegación y Lectura de Códigos QR", level=2)
doc.add_paragraph(
    "1. En la pantalla principal, seleccione la pestaña 'Mis OT' o presione el botón flotante 'Escanear QR'.\n"
    "2. Apunte la cámara del celular al código QR ubicado en el poste SOS o armario de CCTV.\n"
    "3. La aplicación abrirá automáticamente la Ficha del Activo desplegando sus 15 atributos e historial de mantenimientos."
)

# Visual Layout Table representation for QR Scanner
t_qr = doc.add_table(rows=1, cols=2)
t_qr.alignment = WD_TABLE_ALIGNMENT.CENTER
c_qr = t_qr.rows[0].cells
c_qr[0].text = "[ PANTALLA MÓVIL: ESCÁNER QR ]\n\n📷 [Cámara Activa de AppSheet]\nFocus: Código QR del Activo\n\nBotón: 'Escanear de Nuevo' / 'Cerrar'"
c_qr[1].text = "[ FICHA DEL ACTIVO DESPLEGADA ]\n\n📌 Código: SOS-002 | Nombre: Poste SOS PR 15+200\n📍 Ubicación: Sutatenza (Calzada Principal)\n🌐 Coordenadas: 4.8123, -73.6541\n\n[ BOTÓN: INICIAR MANTENIMIENTO ]"

doc.add_paragraph() # Spacing

doc.add_heading("2.3 Paso 3: Diligenciamiento de Checklist Dinámico y Evidencias", level=2)
doc.add_paragraph(
    "1. Al presionar 'Iniciar Mantenimiento', la app cargará automáticamente el checklist correspondiente al tipo de activo (ej. FRM_SOS, FRM_CCTV, FRM_UPS).\n"
    "2. Responda los ítems de inspección (Conforme, No Conforme, No Aplica).\n"
    "3. Adjunte hasta 6 fotografías de evidencia (la app las comprimirá a 600px automáticamente).\n"
    "4. Capture la firma manuscrita digital en la pantalla táctil."
)

doc.add_heading("2.4 Paso 4: Geofencing GPS y Cierre de Orden", level=2)
doc.add_paragraph(
    "1. Al presionar 'Guardar', AppSheet tomará automáticamente las Coordenadas_Cierre mediante la función HERE() y la precisión satelital mediante USERLOCATIONACCURACY().\n"
    "2. Regla de Validación: La app evaluará que la distancia al activo sea menor o igual a 1.0 km:\n"
    "   DISTANCE([Coordenadas_Cierre], LATLONG([ActivoID].[Latitud], [ActivoID].[Longitud])) <= 1.0\n"
    "3. Si se encuentra a más de 1.0 km, el sistema bloqueará el guardado indicando el error."
)

# Box for Bypass Protocol
p_box = doc.add_paragraph()
run_box = p_box.add_run("⚠️ PROTOCOLO DE BYPASS GPS EN TÚNELES O SOMBRA SATELITAL:\nSi el mantenimiento se realiza dentro de un túnel sin señal GPS, el técnico debe registrar las coordenadas del portal de entrada e ingresar en la casilla Observaciones: 'Mantenimiento en túnel - Coordenada de portal registrada por falta de cobertura GPS'.")
run_box.font.size = Pt(10)
run_box.font.bold = True
run_box.font.color.rgb = RGBColor(153, 102, 0)

doc.add_paragraph() # Spacing

# 3. GUÍA ILUSTRADA PARA SUPERVISORES CCO
doc.add_heading("3. Guía Ilustrada para Supervisores CCO (Portal Web)", level=1)

doc.add_heading("3.1 Paso 1: Asignación de Órdenes de Trabajo", level=2)
doc.add_paragraph(
    "1. Ingrese desde su navegador web al portal SGMC en AppSheet.\n"
    "2. Abra la vista 'Órdenes de Trabajo' y haga clic en '+ Agregar OT'.\n"
    "3. Seleccione el Activo, el Técnico Asignado y la Fecha Programada.\n"
    "4. El Automation Bot enviará un correo electrónico de notificación al técnico automáticamente."
)

doc.add_heading("3.2 Paso 2: Monitoreo en Tablero KPI y Alertas de Servicio", level=2)
doc.add_paragraph(
    "1. Ingrese a la vista 'Tablero KPI' para monitorear el porcentaje de cumplimiento y disponibilidad de activos en tiempo real.\n"
    "2. Cuando un técnico marque un activo como 'Fuera de Servicio', la bandeja del CCO recibirá una alerta por correo con la ficha del daño y el informe PDF adjunto."
)

# 4. MATRIZ DE RESOLUCIÓN DE PROBLEMAS
doc.add_heading("4. Matriz de Solución de Problemas (Troubleshooting)", level=1)

t_err = doc.add_table(rows=1, cols=3)
t_err.alignment = WD_TABLE_ALIGNMENT.CENTER
h_e = t_err.rows[0].cells
h_e[0].text = "Problema Observado"
h_e[1].text = "Causa Probable"
h_e[2].text = "Solución Paso a Paso"

errors_data = [
    ("Error: 'Ubicación fuera de rango'", "Técnico a más de 1 km del activo", "Acercarse al activo o aplicar el protocolo de Bypass si es túnel."),
    ("Error: 'Invalid value en GPS'", "Ubicación GPS del celular apagada", "Activar el GPS en los ajustes del smartphone."),
    ("No aparecen activos de otra zona", "Security Filter por SedeID activo", "Verificar en USR_Usuarios que la SedeID asignada sea correcta."),
    ("Lentitud en envío de fotografías", "Señal móvil débil en montaña", "Mantener activado el Background Sync; la app enviará las fotos en segundo plano.")
]

for e in errors_data:
    c = t_err.add_row().cells
    c[0].text = e[0]
    c[1].text = e[1]
    c[2].text = e[2]

import os
RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc.save(os.path.join(RAIZ, "Manuales", "Manual_de_Usuario_SGMC_Con_Diagramas.docx"))
print("Manual_de_Usuario_SGMC_Con_Diagramas.docx generado exitosamente!")
