# -*- coding: utf-8 -*-
"""Utilidades de composicion Word compartidas por los generadores del SGMC."""
import os
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TINTA = RGBColor(0x1C, 0x25, 0x30)
GRIS = RGBColor(0x6E, 0x7A, 0x86)
AZUL = RGBColor(0x1F, 0x5B, 0x99)
ROJO = RGBColor(0xA8, 0x2D, 0x2D)
AMBAR = RGBColor(0xB0, 0x70, 0x0C)
VERDE = RGBColor(0x20, 0x74, 0x54)

# Carpeta de figuras. El generador que importa este modulo la asigna:
#     import _helpers_docx as H;  H.IMG = ruta
IMG = ""

# --------------------------------------------------------------- utilidades
def sombrear(celda, hex_color):
    tc = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def bordes(celda, color="CBD3DB", sz=4, lados=("top", "left", "bottom", "right")):
    tc = celda._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for lado in lados:
        e = OxmlElement(f"w:{lado}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        b.append(e)
    tc.append(b)


def barra_izquierda(celda, color):
    tc = celda._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "18")
    e.set(qn("w:color"), color)
    b.append(e)
    tc.append(b)


def p(doc, texto="", size=10.5, bold=False, color=TINTA, space_after=6,
      space_before=0, align=None, italic=False, izq=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.left_indent = Cm(izq)
    if align:
        par.alignment = align
    if texto:
        r = par.add_run(texto)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        r.font.name = "Segoe UI"
    return par


def titulo(doc, texto, nivel=1, color=TINTA, size=None):
    tam = size or (16 if nivel == 1 else 13 if nivel == 2 else 11.5)
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(18 if nivel == 1 else 14)
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.keep_with_next = True
    r = par.add_run(texto)
    r.font.size = Pt(tam)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Segoe UI"
    return par


def vinneta(doc, texto, size=10.5, color=TINTA, izq=0.8):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.left_indent = Cm(izq)
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(texto)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Segoe UI"
    return par


def bloque(doc, etiqueta, texto, color_hex, color_txt, fondo_hex):
    """Bloque con barra de color a la izquierda."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, fondo_hex)
    barra_izquierda(c, color_hex)
    c.paragraphs[0].text = ""
    par = c.paragraphs[0]
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(6)
    if etiqueta:
        r = par.add_run(etiqueta + "  ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = color_txt
        r.font.name = "Segoe UI"
    r = par.add_run(texto)
    r.font.size = Pt(10)
    r.font.color.rgb = TINTA
    r.font.name = "Segoe UI"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def casilla_respuesta(doc, alto_lineas=5, etiqueta="RESPUESTA"):
    """Caja editable para que el funcional escriba."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, "FFFDF5")
    bordes(c, color="B0700C", sz=8)
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(etiqueta)
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = AMBAR
    r.font.name = "Segoe UI"
    for _ in range(alto_lineas):
        pp = c.add_paragraph()
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.space_before = Pt(0)
        rr = pp.add_run(" ")
        rr.font.size = Pt(10.5)
        rr.font.name = "Segoe UI"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def tabla(doc, encabezados, filas, anchos=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        sombrear(hdr[i], "1F5B99")
        par = hdr[i].paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.space_before = Pt(2)
        r = par.add_run(h)
        r.font.size = Pt(size)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Segoe UI"
    for k, fila in enumerate(filas):
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            if k % 2 == 1:
                sombrear(celdas[i], "F4F6F8")
            par = celdas[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.space_before = Pt(2)
            r = par.add_run(str(v))
            r.font.size = Pt(size)
            r.font.name = "Segoe UI"
            r.font.color.rgb = TINTA
    if anchos:
        for i, w in enumerate(anchos):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def figura(doc, archivo, pie, ancho=16.5):
    ruta = os.path.join(IMG, archivo)
    if not os.path.exists(ruta):
        print("  AVISO: falta la figura", archivo)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(3)
    par.add_run().add_picture(ruta, width=Cm(ancho))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(pie)
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = GRIS
    r.font.name = "Segoe UI"


def pie_de_pagina(doc):
    sec = doc.sections[0]
    par = sec.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("SGMC | Definición funcional y mesa de trabajo | Concesión Transversal del Sisga S.A.S. | Página ")
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS
    r.font.name = "Segoe UI"
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    par._p.append(fld)


