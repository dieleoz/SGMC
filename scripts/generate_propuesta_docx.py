# -*- coding: utf-8 -*-
"""Genera la propuesta de arquitectura y funcionalidad del SGMC, para Direccion
y para el lider funcional.

Salida: entregables/Propuesta_Arquitectura_SGMC.docx

No es un informe de avance: es la propuesta de que se va a construir, por que, y
que decisiones de costo y de propiedad hay que tomar. Valida el diseno desde seis
roles distintos.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

import _helpers_docx as H
from _helpers_docx import (TINTA, GRIS, AZUL, ROJO, AMBAR, VERDE,
                           p, titulo, vinneta, bloque, tabla, figura, pie_de_pagina)
from modelo_objetivo import MODELO, REGLAS, RETIRADAS

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
H.IMG = os.path.join(RAIZ, "docs", "images")
SALIDA = os.path.join(RAIZ, "entregables", "Propuesta_Arquitectura_SGMC.docx")

NUEVAS = [t for t, d in MODELO.items() if d.get("nueva")]
N_COLS = sum(len(d["columnas"]) for d in MODELO.values())
N_REFS = sum(1 for d in MODELO.values() for c in d["columnas"] if c["tipo"] == "Ref")


def construir():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    est = doc.styles["Normal"]
    est.font.name = "Segoe UI"; est.font.size = Pt(10.5)
    est.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    pie_de_pagina(doc)

    # -------------------------------------------------------------- PORTADA
    p(doc, "CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.", size=10, bold=True, color=AZUL, space_after=2)
    p(doc, "Sistema de Gestión de Mantenimiento en Campo", size=10, color=GRIS, space_after=26)
    p(doc, "Propuesta de arquitectura", size=30, bold=True, color=TINTA, space_after=0)
    p(doc, "y funcionalidad", size=30, bold=True, color=AZUL, space_after=16)
    p(doc, "Qué vamos a construir, por qué, y qué hay que decidir", size=13, color=GRIS, space_after=24)

    tabla(doc, ["", ""], [
        ["Dirigido a", "Dirección y liderazgo funcional"],
        ["Fecha", "6 de agosto de 2026"],
        ["Estado", "Propuesta para aprobación"],
        ["Qué se propone", f"Rediseño del modelo de datos: {len(MODELO)} tablas, {N_REFS} referencias, {len(REGLAS)} reglas"],
        ["Qué se necesita", "Tres decisiones de Dirección, en la sección 7"],
    ], anchos=[4.5, 12.0], size=10)

    p(doc, space_after=14)
    bloque(doc, "Por qué esta propuesta y no un informe de avance.",
           "Al auditar el sistema construido se encontró que las relaciones entre sus tablas no "
           "existen: están escritas en la documentación pero no configuradas en la aplicación. Eso "
           "explica por qué, tras un mes, no hay un solo mantenimiento registrado. No es un ajuste "
           "pendiente: hay que rehacer el modelo de datos. Este documento propone cómo.",
           "A82D2D", ROJO, "FBE9E9")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------------ 1. EL DOLOR
    titulo(doc, "1. El problema que el sistema debe resolver", 1)
    p(doc, "Antes de hablar de tablas conviene fijar para qué existe el sistema, porque todo lo "
           "demás se justifica o se descarta contra eso.")
    bloque(doc, "",
           "Garantizar que el mantenimiento se hizo, que quien lo hizo estuvo físicamente frente al "
           "equipo, y que la evidencia que lo respalda es difícil de falsificar.",
           "1F5B99", AZUL, "E2ECF6")
    p(doc, "Todo lo demás — formularios, reportes, tableros — sirve a eso o sobra.", space_before=6)

    titulo(doc, "1.1 Cómo se garantiza la presencia", 2)
    p(doc, "Ningún mecanismo aislado lo consigue. El código QR, por sí solo, no prueba presencia: "
           "una fotografía del QR se escanea desde cualquier parte. El GPS por sí solo tampoco: "
           "falla en túnel y admite manipulación. Lo que funciona es encadenar evidencias "
           "independientes que tendrían que falsificarse todas a la vez y de forma coherente.")
    tabla(doc, ["Eslabón", "Qué aporta"], [
        ["Escaneo del QR físico del equipo",
         "Vincula el registro a ese equipo concreto. Se guarda la hora y la coordenada del momento del escaneo"],
        ["Fotografías tomadas con la cámara de la aplicación",
         "Cada fotografía guarda su propia coordenada, precisión y hora. No se permite elegir de la galería"],
        ["Coordenada de cierre dentro del radio del activo",
         "Impide cerrar el mantenimiento lejos del equipo"],
        ["Marca de tiempo del servidor",
         "La hora la pone el sistema, no el reloj del teléfono, que el usuario puede cambiar"],
        ["Contraste entre escaneo y cierre",
         "Si escaneó en un punto y cerró en otro distante, queda señalado en el reporte de excepciones"],
    ], anchos=[6.0, 10.5])
    bloque(doc, "Una precisión honesta.",
           "Esta cadena no es infalsificable, y conviene decirlo. Lo que hace es elevar el costo de "
           "falsificar por encima del costo de hacer el trabajo. Ese es el objetivo realista de un "
           "control de este tipo, y se consigue.",
           "B0700C", AMBAR, "FDF3DE")
    p(doc, "Un detalle técnico con consecuencia directa: al comprimir las fotografías para no agotar "
           "el plan de datos del técnico, se descartan los metadatos internos de la imagen. Por eso "
           "la fecha, la hora y la coordenada se guardan como datos propios de cada fotografía y no "
           "se confían al archivo. Sin esa decisión, la evidencia no sería defendible.", space_before=6)

    # ------------------------------------------------- 2. INGENIERIA DE MTTO
    titulo(doc, "2. Lo que faltaba: ingeniería de mantenimiento", 1)
    p(doc, "El sistema tal como estaba planteado registraba mantenimientos, pero no los gestionaba. "
           "La diferencia no es semántica: sin plan, alguien tiene que acordarse de crear cada orden "
           "a mano, y lo que no se crea no se hace ni se mide.")
    p(doc, "Se incorporan tres piezas que no existían:")
    tabla(doc, ["Pieza", "Qué habilita"], [
        ["Plan de mantenimiento",
         "Define qué activo se interviene y cada cuánto. De ahí salen las órdenes automáticamente, en "
         "lugar de crearse una por una"],
        ["Taxonomía de modos de falla",
         "Clasifica qué falló y en qué componente. Sin esto no hay tiempo medio entre fallas, ni se "
         "puede saber qué componente falla más, ni pasar algún día de correctivo a predictivo"],
        ["Mediciones con rango en el checklist",
         "Las preguntas numéricas ya guardan valor, unidad y límites. Acumuladas en el tiempo permiten "
         "ver la degradación de un equipo antes de que falle"],
    ], anchos=[5.0, 11.5])
    p(doc, "Con esas tres piezas el sistema pasa de ser un formulario digital a soportar un programa "
           "de mantenimiento preventivo y correctivo, y deja abierta la puerta al predictivo sin "
           "rediseñarlo.", space_before=4)

    # ----------------------------------------------------------- 3. EL MODELO
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "3. El modelo propuesto", 1)
    p(doc, f"{len(MODELO)} tablas, {N_COLS} columnas, {N_REFS} referencias y {len(REGLAS)} reglas. "
           f"El diseño está codificado y se valida automáticamente: hoy pasa las trece "
           f"comprobaciones sin errores.")

    titulo(doc, "3.1 Tablas nuevas", 2)
    tabla(doc, ["Tabla", "Por qué se añade"],
          [[t, MODELO[t]["proposito"].split(".")[0] + "."] for t in NUEVAS],
          anchos=[5.0, 11.5])

    titulo(doc, "3.2 Lo que se retira", 2)
    p(doc, "Cinco tablas y trece campos. En todos los casos por duplicación: el mismo dato guardado "
           "en dos sitios que acaban discrepando.")
    tabla(doc, ["Se retira", "Motivo"],
          [[t, m] for t, m in RETIRADAS.items()], anchos=[5.0, 11.5])

    # --------------------------------------------- 4. VALIDACION POR ROLES
    titulo(doc, "4. Validación del diseño desde seis roles", 1)
    p(doc, "Un diseño puede ser coherente consigo mismo y aun así ser inviable. Se revisó desde seis "
           "puntos de vista distintos; cada uno encontró cosas que los demás no ven.")

    for rol, hallazgos in [
        ("Técnico en campo", [
            ("La app debe devolverle algo, no solo pedirle", "Ve su historial y sus órdenes del día", "Incluido"),
            ("Perder el trabajo por una interrupción lo devuelve al papel", "Borrador local; la inspección queda en curso", "Incluido"),
            ("El plan de datos lo paga él", "Fotografías comprimidas y descarga limitada a su zona", "Incluido"),
            ("El GPS falla en túnel y cañón", "Cierre con excepción justificada, revisado por el supervisor", "Incluido"),
        ]),
        ("Operador del sistema", [
            ("Dar de alta usuarios y activos sin tocar el modelo", "Todo por catálogo", "Incluido"),
            ("Corregir un formulario sin romper el histórico", "Los formularios se versionan y la respuesta guarda su versión", "Incluido"),
            ("Enterarse de que un bot no envió el correo", "Sin resolver en la plataforma", "Pendiente"),
        ]),
        ("Dirección", [
            ("Costo mensual y qué pasa si se deja de pagar", "Ver sección 6", "Decisión"),
            ("De quién son la app, el backend y las fotografías", "Hoy de una cuenta personal. Ver sección 7", "Decisión"),
            ("Qué decisión permite tomar que hoy no se pueda", "Cumplimiento real por zona y por técnico, y disponibilidad por activo", "Incluido"),
        ]),
        ("Interventoría y contrato", [
            ("La evidencia debe ser defendible", "Fotografía con coordenada, hora de servidor, firma y escaneo del QR", "Incluido"),
            ("La cifra debe coincidir con la definición contractual", "Disponibilidad medida por tiempo fuera de servicio", "Confirmar"),
            ("Reconstruir cómo era un formulario en una fecha pasada", "Versionado de formularios", "Incluido"),
        ]),
        ("Capacidad", [
            ("Cuándo se degrada el sistema", "El detalle de checklist multiplica por quince cada mantenimiento", "Vigilar"),
            ("Dónde se guardan las fotografías", "En el Drive del propietario del documento. Ver sección 7", "Bloqueante"),
            ("Retención de cinco años", "Con el inventario completo, el almacenamiento se agota antes", "Bloqueante"),
        ]),
        ("Evolución", [
            ("Qué se rompe si se duplican los activos", "Nada en el modelo; sí en el rendimiento de sincronización", "Vigilar"),
            ("Qué cambio obligaría a migrar datos ya cargados", "El modelo de evidencia fotográfica. Por eso se decide ahora", "Decidido"),
        ]),
    ]:
        titulo(doc, f"4.{['Técnico en campo','Operador del sistema','Dirección','Interventoría y contrato','Capacidad','Evolución'].index(rol)+1} {rol}", 2)
        tabla(doc, ["Lo que pregunta este rol", "Respuesta del diseño", "Estado"],
              [[a, b, c] for a, b, c in hallazgos], anchos=[5.6, 8.4, 2.5], size=9)

    # ------------------------------------------------------- 5. CAPACIDAD
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "5. Hasta dónde aguanta el sistema", 1)
    p(doc, "Cifras calculadas, no estimadas a ojo. Supone cuatro fotografías por mantenimiento, "
           "quince preguntas por checklist y retención de cinco años.")
    tabla(doc, ["Escenario", "Mantenimientos al año", "Almacenamiento a 5 años", "Cuota de 15 GB", "Situación"], [
        ["34 activos, hoy", "489", "1,25 GB", "8 %", "Holgado"],
        ["150 activos", "2.160", "5,54 GB", "37 %", "Vigilar la sincronización"],
        ["500 activos, corredor completo", "7.200", "18,45 GB", "123 %", "No cabe"],
    ], anchos=[5.0, 3.2, 3.4, 2.4, 2.5], size=9)
    p(doc, "Dos conclusiones que conviene retener:", space_before=6)
    vinneta(doc, "El límite de la hoja de cálculo no es la restricción. Se degrada antes la "
                 "sincronización con los celulares, y la tabla que primero llega al umbral es el "
                 "detalle del checklist.")
    vinneta(doc, "La restricción real es el almacenamiento de fotografías, y no por el volumen "
                 "sino por dónde vive. Ver la sección 7.")

    # ------------------------------------------------------------ 6. COSTO
    titulo(doc, "6. Costo y restricciones de la plataforma", 1)
    bloque(doc, "Dato que cambia el alcance.",
           "En el plan gratuito los procesos programados no se ejecutan. Es decir: la generación "
           "automática de órdenes desde el plan de mantenimiento, el correo semanal con las tareas "
           "de cada técnico y el marcado automático de órdenes vencidas requieren plan pagado. Sin "
           "él, el sistema vuelve a depender de que alguien cree cada orden a mano.",
           "B0700C", AMBAR, "FDF3DE")
    tabla(doc, ["Concepto", "Situación"], [
        ["Modalidad de cobro", "Por usuario activo al mes"],
        ["Usuarios registrados hoy", "11, de los cuales 2 inactivos"],
        ["Qué habilita el plan pagado", "Procesos programados, correos automáticos y acceso por interfaz de programación"],
        ["Qué pasa si se deja de pagar", "Los datos permanecen en Google Sheets, propiedad de la Concesión. Se pierde el acceso desde la aplicación, no la información"],
    ], anchos=[5.0, 11.5])
    p(doc, "El presupuesto declarado en el plan original fue de 100 dólares mensuales. Conviene "
           "confirmar la tarifa vigente y el número de usuarios en régimen antes de comprometerla.",
      space_before=4, italic=True, color=GRIS)

    # -------------------------------------------------------- 7. DECISIONES
    titulo(doc, "7. Lo que necesitamos que Dirección decida", 1)
    p(doc, "Tres decisiones. Las dos primeras condicionan la salida a producción.")

    for n, t, cuerpo, riesgo in [
        ("D-A", "Propiedad del backend y del almacenamiento",
         "El documento de datos y las fotografías pertenecen hoy a una cuenta personal de correo, la "
         "del desarrollador. Las imágenes consumen la cuota de esa cuenta, compartida con su correo "
         "personal, con un tope de 15 GB. Se propone trasladar la propiedad a una cuenta corporativa "
         "de la Concesión antes de cargar datos reales.",
         "Con el inventario completo la cuota se agota antes de cumplir la retención de cinco años. "
         "Y mientras la propiedad no se traslade, la continuidad del sistema depende de una persona."),
        ("D-B", "Plan de licenciamiento",
         "Confirmar el paso a plan pagado y el número de usuarios. Sin él no hay generación "
         "automática de órdenes ni notificaciones, que es justamente lo que convierte el sistema en "
         "gestión de mantenimiento.",
         "Sin plan pagado el sistema queda como registro manual, que es la mitad del valor."),
        ("D-C", "Definición contractual de disponibilidad",
         "Se propone medirla por tiempo fuera de servicio sobre el periodo. Si existe una definición "
         "en el contrato o una exigencia de interventoría, debe primar.",
         "Si la cifra que reporta el sistema no coincide con la definición contractual, el primer "
         "informe a interventoría abre una discusión evitable."),
    ]:
        titulo(doc, f"{n}. {t}", 2, color=AZUL)
        p(doc, cuerpo)
        p(doc, "Riesgo si no se decide: " + riesgo, size=9.5, italic=True, color=ROJO)

    # ---------------------------------------------------------- 8. SIGUIENTE
    titulo(doc, "8. Qué sigue", 1)
    for i, (t, s_) in enumerate([
        ("Aprobación de esta propuesta", "Y de las tres decisiones de la sección 7."),
        ("Copia de respaldo", "De la aplicación y del backend, antes de tocar nada."),
        ("Reconstrucción del modelo", "Cablear las relaciones, que es el defecto de raíz, y aplicar el modelo propuesto."),
        ("Carga de datos de prueba", "Poblar el sistema y ejercitarlo. Es lo que revela qué funciona de verdad y qué tablas sobran."),
        ("Entrega para validación", "Sistema completo, con manual de uso por rol, para que el equipo funcional lo opere y corrija sobre algo concreto."),
        ("Piloto en vía", "Con los diez técnicos, solo después de que el paso anterior haya funcionado."),
    ], 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(4)
        r = par.add_run(f"{i}. {t}. ")
        r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = AZUL; r.font.name = "Segoe UI"
        r = par.add_run(s_)
        r.font.size = Pt(10.5); r.font.color.rgb = TINTA; r.font.name = "Segoe UI"

    p(doc, space_after=10)
    bloque(doc, "Sobre el método.",
           "Las definiciones funcionales que no se han podido cerrar se adoptan como supuestos "
           "declarados por escrito, y se construye sobre ellos. Es más rápido corregir una "
           "suposición probada en campo que esperar una definición en abstracto. Todos los supuestos "
           "quedan registrados y cada uno indica qué cuesta cambiarlo.",
           "1F5B99", AZUL, "E2ECF6")

    p(doc, space_after=18)
    t = doc.add_table(rows=1, cols=2)
    for i, et in enumerate(["Por la Dirección", "Por el equipo técnico"]):
        c = t.cell(0, i)
        c.paragraphs[0].text = ""
        pp = c.paragraphs[0]
        r = pp.add_run("\n\n____________________________________")
        r.font.size = Pt(10.5); r.font.name = "Segoe UI"; r.font.color.rgb = TINTA
        pp2 = c.add_paragraph()
        r = pp2.add_run(et + "\nNombre y fecha")
        r.font.size = Pt(9); r.font.color.rgb = GRIS; r.font.name = "Segoe UI"

    doc.save(SALIDA)
    print("Documento generado:", SALIDA)


if __name__ == "__main__":
    construir()
