# -*- coding: utf-8 -*-
"""Genera el documento Word de especificaciones tecnicas As-Built del SGMC.

Salida: entregables/Especificaciones_Tecnicas_SGMC_AsBuilt.docx
Figuras: docs/images/fig_0*.png (generadas por generate_figuras.py)
Utilidades: scripts/_helpers_docx.py

Sustituye la version 1.0, que tenia 15 parrafos, declaraba 17 tablas cuando el
modelo tiene 24, traia la formula de geofencing sobre columnas Latitud/Longitud
que no existen, y marcaba los 16 requerimientos como implementados sin respaldo.

Este documento responde tres preguntas en orden: que hace el sistema, que
ofrece a cada rol, y como funciona. El detalle tecnico va despues, no antes.
El estado de cada requerimiento es el verificado el 2026-08-06 leyendo el
archivo maestro, no el reportado.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import _helpers_docx as H
from _helpers_docx import (TINTA, GRIS, AZUL, ROJO, AMBAR, VERDE,
                           sombrear, bordes, p, titulo, vinneta, bloque,
                           tabla, figura, pie_de_pagina)

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
H.IMG = os.path.join(RAIZ, "docs", "images")
SALIDA = os.path.join(RAIZ, "entregables", "Especificaciones_Tecnicas_SGMC_AsBuilt.docx")


def codigo(doc, texto):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    sombrear(c, "F4F6F8")
    bordes(c, color="CBD3DB", sz=4)
    par = c.paragraphs[0]
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    r = par.add_run(texto)
    r.font.size = Pt(9.5)
    r.font.name = "Consolas"
    r.font.color.rgb = TINTA
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# Estado verificado el 2026-08-06 leyendo BD/Modelo de Datos (2).xlsx
RF = [
 ("RF-001", "Autenticación con la cuenta corporativa",
  "Inicio de sesión con el correo institucional. La aplicación resuelve el usuario con USEREMAIL().",
  "Conforme"),
 ("RF-002", "Operación sin señal celular",
  "La aplicación guarda en el celular y sincroniza al recuperar red. Es capacidad nativa de la plataforma.",
  "Sin verificar"),
 ("RF-003", "Asignación de órdenes de trabajo",
  "El supervisor programa la orden y el sistema la distribuye al técnico. Hay 6 órdenes registradas.",
  "Conforme"),
 ("RF-004", "Restricción de datos por zona",
  "Cada técnico descarga solo los activos de su zona. Hoy usuarios y activos están en sedes disjuntas: "
  "la regla devolvería cero activos.",
  "No conforme"),
 ("RF-005", "Ficha completa del activo",
  "Los 14 atributos de ACT_Activos, poblados en los 34 activos, incluida la ubicación.",
  "Conforme"),
 ("RF-006", "Lectura de código QR",
  "Apertura del activo enfocando la etiqueta. Los 34 activos tienen código QR asignado; la lectura en "
  "campo no se ha probado.",
  "Sin verificar"),
 ("RF-007", "Checklist según el tipo de activo",
  "La aplicación debe abrir el formulario del tipo escaneado. La columna que enlaza tipo y formulario "
  "está vacía en los 18 tipos.",
  "No conforme"),
 ("RF-008", "Formulario de inspección de postes SOS",
  "15 preguntas con secciones, tipos de respuesta, rangos y unidades, construidas en el motor.",
  "Conforme"),
 ("RF-009", "Formularios de CCTV y paneles",
  "Existen hojas planas de plantilla, pero sin preguntas cargadas en el motor de formularios.",
  "No conforme"),
 ("RF-010", "Evidencia fotográfica",
  "Hasta 6 fotografías comprimidas. El modelo las soporta de dos formas incompatibles y los campos "
  "internos solo admiten 2.",
  "No conforme"),
 ("RF-011", "Registro de precisión del GPS",
  "La columna Precision_GPS existe en la tabla de mantenimientos. Nunca se ha poblado.",
  "Parcial"),
 ("RF-012", "Validación de cercanía al activo",
  "La expresión es correcta contra el modelo, pero los 34 activos comparten una coordenada situada en "
  "Bogotá: la regla es inoperante.",
  "No conforme"),
 ("RF-013", "Firma manuscrita",
  "Modelada dos veces, en el mantenimiento y en tabla aparte. La tabla de firmas está vacía.",
  "Parcial"),
 ("RF-014", "Portal web para el centro de control",
  "Vistas web de programación, seguimiento y administración, publicadas en la aplicación.",
  "Conforme"),
 ("RF-015", "Tablero de indicadores",
  "No hay indicador con fórmula, periodicidad ni destinatario acordados.",
  "No conforme"),
 ("RF-016", "Alerta automática por correo",
  "El bot está previsto sobre el cambio de estado a fuera de servicio. Nunca se ha disparado porque no "
  "hay mantenimientos registrados.",
  "Sin verificar"),
]

TABLAS = [
 ("Catálogos y soporte", [
   ("USR_Usuarios", 11, "Usuarios, cargo, rol y sede. Enlaza la sesión con el perfil"),
   ("ROL_Roles", 4, "Administrador, Supervisor, Técnico y Consulta"),
   ("SED_Sedes", 4, "Sedes y unidades funcionales del corredor"),
   ("TIP_TiposActivo", 7, "Los 18 tipos de activo y su formulario de inspección"),
   ("EST_Activo", 2, "Operativo, en mantenimiento, fuera de servicio, retirado"),
   ("FRE_Frecuencias", 4, "Periodicidad del mantenimiento preventivo"),
   ("CAL_Calzadas", 2, "Calzadas del corredor"),
   ("SEN_Sentidos", 2, "Sentidos de circulación"),
   ("FRM_Formularios", 6, "Registro maestro de los 18 formularios")]),
 ("Maestras y checklists", [
   ("ACT_Activos", 14, "Inventario. Código, tipo, sede, PR, calzada, sentido, ubicación, QR, frecuencia"),
   ("CHK_Checklists", 9, "Encabezado de cada inspección ejecutada"),
   ("CHD_ChecklistDetalle", 6, "Resultado ítem por ítem de la inspección")]),
 ("Transaccionales y evidencias", [
   ("OT_OrdenesTrabajo", 12, "Orden programada. Su clave es Numero_OT, no OTID"),
   ("MAN_Mantenimientos", 24, "Ejecución en campo, con Coordenadas_Cierre y Precision_GPS"),
   ("FOT_Fotografias", 5, "Fotografías asociadas al mantenimiento"),
   ("FIR_Firmas", 4, "Firmas del técnico y del supervisor"),
   ("GPS", 8, "Traza de posición para auditoría")]),
 ("Motor de formularios", [
   ("FRM_Preguntas", 17, "Banco único de preguntas, con rangos, unidades y obligatoriedad"),
   ("FRM_Secciones", 4, "Agrupación de preguntas dentro del formulario"),
   ("TPR_TiposRespuesta", 2, "Tipo de dato esperado en cada respuesta"),
   ("LST_ValoresLista", 5, "Opciones de las preguntas de lista"),
   ("FRM_SOS", 11, "Plantilla plana de postes SOS, arquitectura paralela"),
   ("FRM_CCTV", 11, "Plantilla plana de cámaras, arquitectura paralela"),
   ("FRM_PMVF", 11, "Plantilla plana de paneles, arquitectura paralela")]),
]


def construir():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

    est = doc.styles["Normal"]
    est.font.name = "Segoe UI"
    est.font.size = Pt(10.5)
    est.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")

    par = sec.footer.paragraphs[0]
    pie_de_pagina(doc)

    # ---------------------------------------------------------- PORTADA
    p(doc, "CONCESIÓN TRANSVERSAL DEL SISGA S.A.S.", size=10, bold=True, color=AZUL, space_after=2)
    p(doc, "Sistema de Gestión de Mantenimiento en Campo", size=10, color=GRIS, space_after=26)
    p(doc, "Qué hace el SGMC", size=30, bold=True, color=TINTA, space_after=0)
    p(doc, "y cómo funciona", size=30, bold=True, color=AZUL, space_after=16)
    p(doc, "Especificaciones técnicas y arquitectura construida", size=13, color=GRIS, space_after=24)

    tabla(doc, ["", ""], [
        ["Plataforma", "Google AppSheet con backend en Google Sheets"],
        ["Aplicación", "SGMC-886843353"],
        ["Modelo de datos", "24 tablas. Archivo maestro: BD/Modelo de Datos (2).xlsx"],
        ["Fecha", "6 de agosto de 2026"],
        ["Versión", "2.0. Sustituye la versión 1.0 de julio de 2026"],
        ["Estado del sistema", "Prototipo funcional. No apto para campo. Ver sección 10"],
    ], anchos=[4.5, 12.0], size=10)

    p(doc, space_after=14)
    bloque(doc, "Sobre el estado que declara este documento.",
           "La versión anterior marcaba los 16 requerimientos como implementados. Este documento declara el estado "
           "verificado el 6 de agosto de 2026 leyendo el archivo de producción, no el reportado. De los 16 "
           "requerimientos, 5 están conformes, 2 parciales, 3 sin verificar y 6 no conformes. El detalle está en la "
           "sección 9 y lo que falta para cerrarlos, en la sección 10.",
           "B0700C", AMBAR, "FDF3DE")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------- 1. QUÉ HACE
    titulo(doc, "1. Qué hace el sistema", 1)
    p(doc, "El SGMC reemplaza el registro en papel del mantenimiento del corredor por un registro digital con "
           "evidencia verificable de que el trabajo se hizo, dónde se hizo y en qué estado quedó el activo.")
    p(doc, "Ataca cuatro problemas concretos de la operación actual:")
    tabla(doc, ["Problema de hoy", "Qué hace el sistema"], [
        ["No hay forma de comprobar que el técnico estuvo frente al activo",
         "Exige la posición del celular al cerrar y la compara con la coordenada registrada del activo. Guarda "
         "además la precisión del satélite"],
        ["Buena parte del corredor no tiene señal celular",
         "El técnico diligencia todo sin red. El registro queda en el celular y sube solo cuando vuelve la señal"],
        ["Cada tipo de activo se inspecciona distinto y los formatos se mezclan",
         "Al escanear el código QR, la aplicación abre el checklist que corresponde a ese tipo de activo"],
        ["El centro de control se entera tarde de una falla",
         "Cuando un activo queda fuera de servicio, envía de inmediato un correo con el informe y las fotografías"],
    ], anchos=[6.0, 10.5])

    # ---------------------------------------------------------- 2. QUÉ OFRECE
    titulo(doc, "2. Qué ofrece a cada rol", 1)
    p(doc, "El mismo sistema se ve distinto según quién entre. La aplicación resuelve el rol a partir del correo "
           "con el que se inicia sesión.")
    tabla(doc, ["Rol", "Qué puede hacer", "Qué gana"], [
        ["Técnico",
         "Ver sus órdenes del día, abrir el activo por QR, responder el checklist, tomar fotografías, firmar y "
         "cerrar. Todo sin señal",
         "Deja de cargar formatos en papel y de transcribirlos después. La evidencia queda tomada en el sitio"],
        ["Supervisor",
         "Programar y asignar órdenes, revisar la evidencia recibida, aprobar o devolver, y consultar el tablero",
         "Sabe qué se hizo y qué no sin llamar a nadie, y aprueba con la evidencia a la vista"],
        ["Administrador",
         "Mantener usuarios, activos, catálogos y los formularios de inspección de cada tipo de activo",
         "Cambia un checklist desde el navegador y llega a los celulares en la siguiente sincronización"],
        ["Consulta",
         "Leer y exportar información, sin modificar nada",
         "Acceso a la información sin riesgo de alterarla"],
    ], anchos=[2.6, 7.4, 6.5])
    figura(doc, "fig_01_actores.png", "Figura 1. Actores del sistema y alcance de cada uno.")

    # ---------------------------------------------------------- 3. QUÉ GESTIONA
    titulo(doc, "3. Qué gestiona", 1)
    p(doc, "34 activos catalogados sobre 18 tipos, agrupados en cuatro categorías:")
    tabla(doc, ["Categoría", "Tipos de activo"], [
        ["ITS", "Postes SOS, cámaras CCTV, paneles de mensaje variable fijo y móvil, sensores meteorológicos y "
                "ambientales, básculas de pesaje"],
        ["Eléctrico", "Generadores, sistemas UPS, subestaciones"],
        ["Comunicaciones", "Fibra óptica"],
        ["Tecnología de la información", "Servidores, NAS, switches, routers, firewalls, videowall"],
    ], anchos=[4.0, 12.5])
    p(doc, "Cada activo lleva su código, tipo, sede, punto de referencia vial, calzada, sentido, coordenada, estado, "
           "frecuencia de mantenimiento y código QR.")

    # ---------------------------------------------------------- 4. CÓMO FUNCIONA
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "4. Cómo funciona", 1)
    p(doc, "La solución se apoya en componentes gestionados. No hay servidores propios, ni compilación de una "
           "aplicación, ni publicación en tiendas: el técnico instala la aplicación AppSheet desde la tienda de su "
           "celular y entra con su cuenta corporativa.")
    figura(doc, "fig_06_arquitectura.png", "Figura 6. Arquitectura de la solución en tres capas.")

    titulo(doc, "4.1 El ciclo del técnico", 2)
    p(doc, "Es el flujo que justifica el sistema. Ocurre en su mayor parte sin conexión.")
    figura(doc, "fig_02_flujo_tecnico.png", "Figura 2. Ciclo del técnico en campo, paso a paso.")

    titulo(doc, "4.2 El ciclo del supervisor", 2)
    p(doc, "El supervisor programa la orden en el portal y la asigna a un técnico. El sistema le asigna número y un "
           "bot notifica por correo al técnico, que la recibe en su celular al sincronizar. Cuando el técnico cierra "
           "y sincroniza, el mantenimiento aparece en el portal con su checklist, fotografías, firma y coordenada de "
           "cierre. El supervisor aprueba o devuelve con observación, y al aprobar la orden se cierra.")
    figura(doc, "fig_04_ciclo_ot.png", "Figura 4. Ciclo de vida de la orden de trabajo.")

    titulo(doc, "4.3 Cómo opera sin señal", 2)
    p(doc, "Al sincronizar con red, la aplicación descarga al celular los activos y órdenes que corresponden al "
           "usuario, y los mantiene disponibles sin conexión. El técnico trabaja contra esa copia local: abre el "
           "activo, responde, fotografía y firma. Al guardar, el registro entra en una cola local con indicador "
           "visible. Cuando el celular recupera señal, la cola sube en segundo plano sin que el técnico tenga que "
           "hacer nada.")
    p(doc, "Las fotografías se comprimen antes de subir, para no agotar el plan de datos del técnico.")

    # ---------------------------------------------------------- 5. MODELO
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "5. Modelo de datos", 1)
    p(doc, "24 tablas en cuatro grupos. La cadena operativa va del activo a la orden, de la orden a la ejecución en "
           "campo, y de la ejecución a sus evidencias.")
    figura(doc, "fig_07_modelo_datos.png", "Figura 7. Las 24 tablas y sus relaciones principales.")

    for grupo, filas in TABLAS:
        titulo(doc, f"5.{[g for g, _ in TABLAS].index(grupo) + 1} {grupo} ({len(filas)})", 2)
        tabla(doc, ["Tabla", "Col.", "Función"],
              [[n, str(c), f] for n, c, f in filas], anchos=[4.4, 1.4, 10.7])

    titulo(doc, "5.5 Advertencias sobre el modelo", 2)
    vinneta(doc, "La nomenclatura de claves no es uniforme. OT_OrdenesTrabajo no tiene columna OTID: su clave es "
                 "Numero_OT, con valores del tipo OT-0001. Otras tablas la referencian como OTID.")
    vinneta(doc, "ACT_Activos guarda un único campo Ubicacion de tipo coordenada. No existen columnas Latitud y "
                 "Longitud separadas. Cualquier expresión que las invoque falla.")
    vinneta(doc, "Las evidencias están modeladas dos veces: como campos dentro de MAN_Mantenimientos y como tablas "
                 "aparte. Hay que elegir una vía.")
    vinneta(doc, "Coexisten dos arquitecturas de formularios: el motor relacional FRM_Preguntas y las hojas planas "
                 "FRM_SOS, FRM_CCTV y FRM_PMVF, con esquemas distintos.")

    # ---------------------------------------------------------- 6. REGLAS
    titulo(doc, "6. Reglas de negocio y expresiones", 1)

    titulo(doc, "6.1 Validación de cercanía al activo", 2)
    p(doc, "Impide cerrar un mantenimiento lejos del activo. Se evalúa sobre el campo Coordenadas_Cierre de "
           "MAN_Mantenimientos.")
    codigo(doc, "DISTANCE([Coordenadas_Cierre], [ActivoID].[Ubicacion]) <= 1.0")
    p(doc, "Mensaje de error, en texto plano:")
    codigo(doc, "Ubicación fuera de rango: debe estar a menos de 1.0 km del activo.")
    bloque(doc, "Corrección respecto a la versión anterior.",
           "La versión 1.0 de este documento traía la expresión sobre [ActivoID].[Latitud] y [ActivoID].[Longitud]. "
           "Esas columnas no existen en el modelo: falla en ejecución. La expresión válida es la de arriba.",
           "A82D2D", ROJO, "FBE9E9")

    titulo(doc, "6.2 Captura de precisión del satélite", 2)
    p(doc, "Valor inicial del campo Precision_GPS. Registra el margen de error en metros, para distinguir un cierre "
           "legítimo de uno con GPS deficiente.")
    codigo(doc, "USERLOCATIONACCURACY()")

    titulo(doc, "6.3 Restricción de datos por zona", 2)
    p(doc, "Filtro de seguridad sobre ACT_Activos y OT_OrdenesTrabajo. Se evalúa en el servidor al iniciar sesión, "
           "de modo que el celular solo descarga lo que corresponde al usuario.")
    codigo(doc, '[SedeID] = LOOKUP(USEREMAIL(), "USR_Usuarios", "Correo", "SedeID")')
    p(doc, "Esta regla no es cosmética: es lo que mantiene el volumen de descarga bajo y la operación offline "
           "viable. Hoy no puede aplicarse porque usuarios y activos están en sedes disjuntas.")

    titulo(doc, "6.4 Alerta por activo fuera de servicio", 2)
    p(doc, "Automatización que reacciona al cambio de estado y envía un correo con el informe de la falla en PDF, "
           "con la ubicación y las fotografías, al centro de control y al supervisor de la zona.")
    codigo(doc, 'Evento:     cambio de datos sobre MAN_Mantenimientos\n'
                'Condición:  [Estado Final] = "Fuera de servicio"\n'
                'Acción:     enviar correo con informe PDF adjunto')

    # ---------------------------------------------------------- 7. SEGURIDAD
    titulo(doc, "7. Seguridad y control de acceso", 1)
    tabla(doc, ["Control", "Cómo opera"], [
        ["Identidad", "Inicio de sesión con la cuenta corporativa. No hay contraseñas propias del sistema"],
        ["Perfil", "El correo de la sesión resuelve el usuario en USR_Usuarios y de ahí su rol y su sede"],
        ["Alcance de datos", "El filtro de seguridad limita lo que se descarga al celular, no solo lo que se ve"],
        ["Permisos por rol", "Cuatro perfiles: Administrador, Supervisor, Técnico y Consulta"],
        ["Trazabilidad", "Cada mantenimiento registra usuario, fecha y hora de registro, y coordenada de cierre"],
    ], anchos=[4.2, 12.3])

    # ---------------------------------------------------------- 8. QUÉ NO HACE
    titulo(doc, "8. Qué no hace el sistema", 1)
    p(doc, "Declarar los límites evita expectativas equivocadas. Hoy el sistema no cubre:")
    vinneta(doc, "Generación automática de órdenes a partir de la frecuencia de mantenimiento del activo. Las "
                 "órdenes se crean manualmente.")
    vinneta(doc, "Gestión de repuestos, inventario de almacén o costos de la intervención.")
    vinneta(doc, "Integración con tableros externos o con mesas de ayuda.")
    vinneta(doc, "Flujo de mantenimiento correctivo iniciado desde el campo sin orden previa.")
    vinneta(doc, "Firma de interventoría o valor probatorio de la firma frente a terceros.")

    # ---------------------------------------------------------- 9. MATRIZ RF
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(doc, "9. Estado de los requerimientos", 1)
    p(doc, "Estado verificado el 6 de agosto de 2026 leyendo el archivo maestro de producción. Conforme significa "
           "comprobado en el archivo. Parcial, que la estructura existe pero el dato no. Sin verificar, que depende "
           "de una prueba en campo que no se ha hecho. No conforme, que hoy no funcionaría.")
    tabla(doc, ["ID", "Requerimiento", "Situación verificada", "Estado"],
          [[i, n, d, e] for i, n, d, e in RF], anchos=[1.6, 4.3, 8.2, 2.4], size=9)
    p(doc, "Resumen: 5 conformes, 2 parciales, 3 sin verificar y 6 no conformes.", bold=True, space_before=4)

    # ---------------------------------------------------------- 10. FALTA
    titulo(doc, "10. Qué falta para llevarlo a campo", 1)
    p(doc, "Ocho hallazgos bloqueantes, verificados contra el archivo. La mayoría no son errores de programación "
           "sino decisiones de operación que aún no se han tomado; por eso cada uno remite a una decisión del "
           "documento de mesa de trabajo.")
    tabla(doc, ["#", "Qué está bloqueado", "Decisión que lo desbloquea"], [
        ["B-01", "Los 34 activos comparten una coordenada situada en Bogotá. La validación de cercanía es "
                 "inoperante hasta levantar las coordenadas reales en campo", "D-01"],
        ["B-02", "La columna que enlaza cada tipo de activo con su formulario está vacía en los 18 tipos", "D-09"],
        ["B-03", "Usuarios y activos están en sedes disjuntas: el filtro dejaría a cada técnico sin activos", "D-03"],
        ["B-04", "Solo 1 de 18 formularios tiene su banco de preguntas construido", "D-09"],
        ["B-05", "El único checklist existente referencia una orden de trabajo que no existe", "D-06"],
        ["B-06", "Fotografías, firmas y posición están modeladas dos veces, de forma incompatible", "D-10"],
        ["B-07", "Ningún mantenimiento se ha ejecutado de extremo a extremo: cuatro tablas están vacías", "Todas"],
        ["B-08", "El detalle de checklist guarda el texto de la pregunta, sin trazabilidad al banco", "D-11"],
    ], anchos=[1.6, 11.5, 3.4])

    bloque(doc, "Conclusión sobre el estado del sistema.",
           "Lo construido es un prototipo funcional con el modelo de datos completo y los catálogos poblados. No es "
           "un sistema listo para campo: el ciclo de mantenimiento nunca se ha ejecutado completo y las dos reglas "
           "que dan valor al sistema, la validación de cercanía y la restricción por zona, hoy no funcionarían con "
           "los datos cargados.",
           "A82D2D", ROJO, "FBE9E9")

    p(doc, "El detalle de cada hallazgo, con la evidencia y el comando para reproducirla, está en el dictamen de "
           "auditoría del 6 de agosto de 2026. Las decisiones D-01 a D-14 están en el documento de definición "
           "funcional y mesa de trabajo.", space_before=6)

    p(doc, space_after=20)
    p(doc, "Especificaciones técnicas As-Built del SGMC. Versión 2.0. Concesión Transversal del Sisga S.A.S.",
      size=9, italic=True, color=GRIS)

    doc.save(SALIDA)
    print("Documento generado:", SALIDA)


if __name__ == "__main__":
    construir()
